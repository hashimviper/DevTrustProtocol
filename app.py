"""
DevTrustProtocol v3.0
Run: python -m streamlit run app.py
"""
import os, re, io, json, time, base64, zipfile, hashlib, pathlib, datetime
import streamlit as st
from dotenv import load_dotenv

# ──────────────────────────────────────────────────────────
# 0. ENV LOADING
# ──────────────────────────────────────────────────────────
_BASE_DIR = pathlib.Path(__file__).parent

def _load_env_file(filename):
    p = _BASE_DIR / filename
    if p.exists():
        load_dotenv(dotenv_path=p, override=True); return True
    return False

_load_env_file("Apikeytop.env")
_load_env_file("chatapikey.env")
load_dotenv(dotenv_path=_BASE_DIR / ".env")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY","")
GROQ_API_KEY   = os.getenv("GROQ_API_KEY","")

# ──────────────────────────────────────────────────────────
# 1. PAGE CONFIG
# ──────────────────────────────────────────────────────────
st.set_page_config(
    page_title="DevTrustProtocol",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={},
)

# ──────────────────────────────────────────────────────────
# 2. MULTI-USER CREDENTIALS (User = developer, Organization = recruiter)
# ──────────────────────────────────────────────────────────
# Users (developers) — 5 accounts
USERS = {
    "alex_dev":     {"password": "dev@2025!",    "role": "User",  "email": "alex@devmail.io",     "display": "Alex Chen"},
    "priya_codes":  {"password": "priya#code9",  "role": "User",  "email": "priya@outlook.com",   "display": "Priya Sharma"},
    "marco_build":  {"password": "marco$3build", "role": "User",  "email": "marco@gmail.com",     "display": "Marco Rossi"},
    "sana_eng":     {"password": "sana!eng42",   "role": "User",  "email": "sana@proton.me",      "display": "Sana Malik"},
    "dev_juno":     {"password": "juno%2025x",   "role": "User",  "email": "juno@devhub.net",     "display": "Juno Park"},
}
# Organizations (recruiters) — 5 accounts
ORGS = {
    "techcorp_hr":   {"password": "tc#hire25",    "role": "Recruiter", "display": "TechCorp Talent",    "email": "talent@techcorp.io"},
    "nexus_recruit": {"password": "nx$rec2025",   "role": "Recruiter", "display": "Nexus Recruitment",  "email": "hire@nexusjobs.com"},
    "buildspace_org":{"password": "bs!org77",     "role": "Recruiter", "display": "BuildSpace Labs",    "email": "careers@buildspace.dev"},
    "openbridge_hr": {"password": "ob#hr2025",    "role": "Recruiter", "display": "OpenBridge HR",      "email": "hr@openbridge.org"},
    "devlink_teams": {"password": "dl@team99",    "role": "Recruiter", "display": "DevLink Teams",      "email": "team@devlink.io"},
}
CREDENTIALS = {**USERS, **ORGS}

ROLE_PERMISSIONS = {
    "User":      {"upload_project":True,"upload_resume":True,"sanitize":True,"score":True,
                  "view_raw_files":True,"view_dashboard":True,"download_resume":True,"view_research":True},
    "Recruiter": {"upload_project":False,"upload_resume":False,"sanitize":False,"score":False,
                  "view_raw_files":False,"view_dashboard":True,"download_resume":True,"view_research":False},
}

def can(permission):
    return ROLE_PERMISSIONS.get(st.session_state.get("role",""),{}).get(permission,False)

# ──────────────────────────────────────────────────────────
# 3. SHARED STATE (multi-user submissions store)
# ──────────────────────────────────────────────────────────
_SHARED_STATE_FILE = _BASE_DIR / "devtrust_shared_state.json"
_NOTIFICATIONS_FILE = _BASE_DIR / "devtrust_notifications.json"

def _load_all_submissions() -> dict:
    try:
        if _SHARED_STATE_FILE.exists():
            with open(_SHARED_STATE_FILE,"r",encoding="utf-8") as f:
                data = json.load(f)
            # Support both old single-entry format and new multi-entry
            if "submissions" in data:
                return data
            elif "result" in data:
                return {"submissions": {"legacy_project": data}}
    except Exception:
        pass
    return {"submissions": {}}

def _save_all_submissions(data: dict):
    try:
        with open(_SHARED_STATE_FILE,"w",encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    except Exception as e:
        st.warning(f"Save error: {e}")

def save_submission(project_name: str, result: dict, profile: dict,
                    resume_text: str, resume_bytes: bytes, resume_ext: str,
                    resume_filename: str, uploaded_by: str,
                    project_type: str = "Software Development",
                    description: str = ""):
    all_data = _load_all_submissions()
    key = re.sub(r"[^a-zA-Z0-9_]","_", project_name.strip() or "project")
    all_data["submissions"][key] = {
        "result": result, "profile": profile, "resume_text": resume_text,
        "resume_bytes_b64": base64.b64encode(resume_bytes).decode() if resume_bytes else None,
        "resume_ext": resume_ext, "resume_filename": resume_filename,
        "uploaded_by": uploaded_by,
        "project_type": project_type,
        "description": description,
        "saved_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    _save_all_submissions(all_data)
    _push_notification("org", f"📥 New project '{project_name}' submitted by {uploaded_by}", project_name)
    _push_notification(uploaded_by, f"✅ Your project '{project_name}' was submitted successfully!")

def get_all_submissions() -> list:
    data = _load_all_submissions()
    items = []
    for key, sub in data.get("submissions",{}).items():
        sub["_key"] = key
        items.append(sub)
    items.sort(key=lambda x: x.get("result",{}).get("trust_score",0), reverse=True)
    return items

def delete_submission(key: str):
    data = _load_all_submissions()
    data["submissions"].pop(key, None)
    _save_all_submissions(data)

def load_shared_state():
    """Legacy compat"""
    subs = get_all_submissions()
    return subs[0] if subs else None

def clear_shared_state():
    _save_all_submissions({"submissions": {}})

# ──────────────────────────────────────────────────────────
# 4. NOTIFICATIONS
# ──────────────────────────────────────────────────────────
def _load_notifications() -> dict:
    try:
        if _NOTIFICATIONS_FILE.exists():
            with open(_NOTIFICATIONS_FILE,"r",encoding="utf-8") as f:
                return json.load(f)
    except Exception: pass
    return {}

def _save_notifications(data: dict):
    try:
        with open(_NOTIFICATIONS_FILE,"w",encoding="utf-8") as f:
            json.dump(data, f, indent=2)
    except Exception: pass

def _push_notification(target: str, message: str, project_key: str = ""):
    data = _load_notifications()
    if target not in data: data[target] = []
    data[target].append({
        "msg": message, "time": time.strftime("%Y-%m-%d %H:%M:%S"),
        "read": False, "project_key": project_key
    })
    _save_notifications(data)

def get_notifications(username: str, role: str) -> list:
    data = _load_notifications()
    notifs = []
    if role == "Recruiter":
        notifs = data.get("org", [])
    else:
        notifs = data.get(username, [])
    return list(reversed(notifs))

def mark_all_read(username: str, role: str):
    data = _load_notifications()
    key = "org" if role == "Recruiter" else username
    for n in data.get(key, []):
        n["read"] = True
    _save_notifications(data)

def unread_count(username: str, role: str) -> int:
    return sum(1 for n in get_notifications(username, role) if not n.get("read"))

# ──────────────────────────────────────────────────────────
# 5. GITHUB SAMPLE PROJECTS (20 real repos, sanitized)
# ──────────────────────────────────────────────────────────
GITHUB_PROJECTS = [
    {"name":"FastAPI Production Boilerplate","user":"gh_user_1","github":"tiangolo/full-stack-fastapi-template",
     "description":"A full-stack FastAPI + React + PostgreSQL production template with Docker, Traefik, and automated tests. Demonstrates async API design, JWT auth, and CI/CD pipelines.","language":"Python","stars":"26k",
     "score":91,"grade":"A","working":94,"recommendation":"Strong Hire","complexity":89,"security":92,"docs":88,"maintainability":90,
     "tech_stack":["FastAPI","Python","React","PostgreSQL","Docker","Redis"],
     "capabilities":["REST API Design","Async Programming","Container Orchestration","JWT Authentication"],
     "red_flags":[],"summary":"Exemplary full-stack project showcasing production-grade FastAPI with complete DevOps pipeline."},
    {"name":"Real-time Chat with WebSockets","user":"gh_user_2","github":"realtime-chat-app",
     "description":"Scalable real-time messaging application built with Node.js, Socket.io, and React. Features rooms, private DMs, typing indicators, read receipts, and file sharing.",
     "language":"JavaScript","stars":"8.4k",
     "score":87,"grade":"A","working":91,"recommendation":"Strong Hire","complexity":84,"security":83,"docs":85,"maintainability":88,
     "tech_stack":["Node.js","Socket.io","React","MongoDB","Redis"],
     "capabilities":["WebSocket Architecture","State Management","Real-time UX"],
     "red_flags":["Minor: Missing rate limiting on message endpoints"],"summary":"Strong real-time architecture with clean separation of concerns."},
    {"name":"ML Pipeline Orchestrator","user":"gh_user_3","github":"ml-pipeline-orchestrator",
     "description":"End-to-end machine learning pipeline framework with experiment tracking, model versioning, and automated deployment. Integrates with MLflow, DVC, and Kubernetes.",
     "language":"Python","stars":"12k",
     "score":93,"grade":"A","working":96,"recommendation":"Strong Hire","complexity":94,"security":90,"docs":95,"maintainability":91,
     "tech_stack":["Python","MLflow","DVC","Kubernetes","Airflow","PyTorch"],
     "capabilities":["MLOps","Pipeline Design","Model Versioning","Kubernetes Deployment"],
     "red_flags":[],"summary":"Outstanding MLOps project with best-in-class tooling and documentation."},
    {"name":"Distributed Task Queue","user":"gh_user_4","github":"distributed-task-queue",
     "description":"High-performance distributed task queue system inspired by Celery. Built from scratch in Go with Redis backend, supporting priorities, retries, dead-letter queues, and web dashboard.",
     "language":"Go","stars":"5.2k",
     "score":88,"grade":"A","working":92,"recommendation":"Hire","complexity":91,"security":87,"docs":82,"maintainability":89,
     "tech_stack":["Go","Redis","gRPC","Docker","Prometheus"],
     "capabilities":["Distributed Systems","Concurrency","Message Queuing","Systems Programming"],
     "red_flags":["Docs could cover failure modes in more depth"],"summary":"Well-engineered distributed system demonstrating strong systems programming skills."},
    {"name":"E-Commerce Microservices","user":"gh_user_5","github":"ecommerce-microservices",
     "description":"Complete microservices-based e-commerce platform with 8 independent services: catalog, cart, orders, payments, notifications, auth, search, and analytics.",
     "language":"TypeScript","stars":"9.1k",
     "score":89,"grade":"A","working":93,"recommendation":"Strong Hire","complexity":92,"security":88,"docs":87,"maintainability":90,
     "tech_stack":["TypeScript","Node.js","Kafka","PostgreSQL","Redis","Docker","K8s"],
     "capabilities":["Microservices Architecture","Event-Driven Design","API Gateway Pattern"],
     "red_flags":[],"summary":"Comprehensive microservices implementation with excellent service isolation."},
    {"name":"React Design System","user":"gh_user_1","github":"react-design-system",
     "description":"Production-ready React component library with 60+ components, full TypeScript support, Storybook documentation, automated visual regression tests, and zero dependencies.",
     "language":"TypeScript","stars":"7.3k",
     "score":85,"grade":"B","working":97,"recommendation":"Hire","complexity":78,"security":82,"docs":94,"maintainability":93,
     "tech_stack":["React","TypeScript","Storybook","Jest","Chromatic","CSS-in-JS"],
     "capabilities":["Component Architecture","Accessibility","TypeScript","Design Systems"],
     "red_flags":["Bundle size could be optimised with tree-shaking"],"summary":"Polished design system with exceptional documentation and test coverage."},
    {"name":"Blockchain Smart Contracts","user":"gh_user_2","github":"defi-smart-contracts",
     "description":"DeFi protocol implementing AMM liquidity pools, yield farming, and governance tokens on Ethereum. Includes formal verification, 100% test coverage, and audit reports.",
     "language":"Solidity","stars":"4.8k",
     "score":86,"grade":"B","working":89,"recommendation":"Hire","complexity":93,"security":91,"docs":80,"maintainability":82,
     "tech_stack":["Solidity","Hardhat","Ethers.js","OpenZeppelin","TypeScript"],
     "capabilities":["Smart Contract Security","DeFi Protocol Design","Formal Verification"],
     "red_flags":["Complex re-entrancy guard logic needs inline comments"],"summary":"Secure smart contract implementation with formal verification — rare skill combination."},
    {"name":"Kubernetes Operator Framework","user":"gh_user_3","github":"k8s-operator-framework",
     "description":"Custom Kubernetes operator framework for managing stateful applications. Implements reconciliation loops, CRD management, and automated failover with comprehensive e2e tests.",
     "language":"Go","stars":"6.5k",
     "score":90,"grade":"A","working":93,"recommendation":"Strong Hire","complexity":95,"security":89,"docs":88,"maintainability":87,
     "tech_stack":["Go","Kubernetes","controller-runtime","Prometheus","Helm"],
     "capabilities":["Kubernetes Internals","Operator Pattern","Cloud Native","Reconciliation"],
     "red_flags":[],"summary":"Expert-level Kubernetes operator with production-ready reliability patterns."},
    {"name":"NLP Text Analytics Engine","user":"gh_user_4","github":"nlp-analytics-engine",
     "description":"Modular NLP pipeline for sentiment analysis, entity recognition, and topic modelling. Supports BERT fine-tuning, multilingual models, and batch processing at scale.",
     "language":"Python","stars":"3.9k",
     "score":82,"grade":"B","working":87,"recommendation":"Hire","complexity":86,"security":78,"docs":84,"maintainability":81,
     "tech_stack":["Python","HuggingFace","PyTorch","FastAPI","Celery","PostgreSQL"],
     "capabilities":["NLP","Model Fine-tuning","API Design","Async Processing"],
     "red_flags":["Inference latency optimisation not implemented","Missing model versioning"],"summary":"Solid NLP implementation with good modularity; could benefit from latency optimisation."},
    {"name":"CI/CD Platform","user":"gh_user_5","github":"cicd-platform",
     "description":"Self-hosted CI/CD platform with parallel pipeline execution, artifact management, secret scanning, and Slack/GitHub integrations. Handles 500+ builds per day in production.",
     "language":"Go","stars":"11k",
     "score":92,"grade":"A","working":95,"recommendation":"Strong Hire","complexity":93,"security":94,"docs":90,"maintainability":92,
     "tech_stack":["Go","PostgreSQL","Docker","gRPC","Redis","Prometheus","Grafana"],
     "capabilities":["DevOps Tooling","Pipeline Architecture","Secret Management","Observability"],
     "red_flags":[],"summary":"Production-grade CI/CD platform with impressive scale and security practices."},
    {"name":"Graph Neural Network Library","user":"gh_user_1","github":"gnn-library",
     "description":"Flexible GNN library supporting GCN, GAT, GraphSAGE, and custom message-passing schemes. Benchmarked on OGB datasets with reproducible results and tutorial notebooks.",
     "language":"Python","stars":"2.1k",
     "score":79,"grade":"B","working":84,"recommendation":"Interview Further","complexity":88,"security":72,"docs":82,"maintainability":76,
     "tech_stack":["Python","PyTorch Geometric","NumPy","NetworkX","Jupyter"],
     "capabilities":["Graph ML","Research Implementation","Benchmarking"],
     "red_flags":["No GPU memory optimisation","Limited distributed training support"],"summary":"Technically strong GNN library with good research quality; production hardening needed."},
    {"name":"Serverless Data Pipeline","user":"gh_user_2","github":"serverless-data-pipeline",
     "description":"Cloud-agnostic serverless data pipeline framework for ETL workloads. Abstracts AWS Lambda, GCP Cloud Functions, and Azure Functions behind a unified SDK.",
     "language":"Python","stars":"4.4k",
     "score":83,"grade":"B","working":88,"recommendation":"Hire","complexity":82,"security":80,"docs":86,"maintainability":84,
     "tech_stack":["Python","AWS Lambda","Terraform","Step Functions","DynamoDB","S3"],
     "capabilities":["Serverless Architecture","Cloud Abstraction","ETL Design","IaC"],
     "red_flags":["Cost estimation tools missing"],"summary":"Practical serverless ETL with good cloud abstraction layer."},
    {"name":"Terminal Emulator in Rust","user":"gh_user_3","github":"terminal-emulator-rust",
     "description":"VT100/VT220-compatible terminal emulator built from scratch in Rust. Implements full ANSI escape code parsing, sixel graphics, and ligature font rendering.",
     "language":"Rust","stars":"3.3k",
     "score":84,"grade":"B","working":86,"recommendation":"Hire","complexity":96,"security":85,"docs":79,"maintainability":80,
     "tech_stack":["Rust","OpenGL","Wayland","X11","wgpu"],
     "capabilities":["Systems Programming","Rust","Terminal Protocols","Graphics"],
     "red_flags":["Documentation for escape code coverage incomplete"],"summary":"Impressive low-level systems work in Rust; deep technical expertise demonstrated."},
    {"name":"Auth Service with MFA","user":"gh_user_4","github":"auth-service-mfa",
     "description":"Enterprise-grade authentication service with TOTP/FIDO2 MFA, OAuth 2.0, SAML 2.0, and SCIM provisioning. Zero-trust architecture with audit logging and anomaly detection.",
     "language":"Python","stars":"5.7k",
     "score":94,"grade":"A","working":97,"recommendation":"Strong Hire","complexity":90,"security":97,"docs":92,"maintainability":93,
     "tech_stack":["Python","FastAPI","PostgreSQL","Redis","FIDO2","SAML","OAuth2"],
     "capabilities":["Security Engineering","Identity Protocols","Zero-Trust","Audit Logging"],
     "red_flags":[],"summary":"Best-in-class security implementation with comprehensive protocol support."},
    {"name":"Vector Database Engine","user":"gh_user_5","github":"vector-db-engine",
     "description":"High-performance vector similarity search engine using HNSW and IVF-PQ indexing. Written in C++ with Python bindings, supporting 100M+ vectors with sub-millisecond queries.",
     "language":"C++","stars":"8.9k",
     "score":91,"grade":"A","working":93,"recommendation":"Strong Hire","complexity":97,"security":86,"docs":88,"maintainability":85,
     "tech_stack":["C++","Python","SIMD","OpenMP","gRPC","Prometheus"],
     "capabilities":["Search Algorithms","High Performance Computing","C++ Systems","SIMD"],
     "red_flags":["Python binding docs lag behind C++ API"],"summary":"Exceptional performance engineering with state-of-the-art similarity search."},
    {"name":"Browser DevTools Extension","user":"gh_user_1","github":"devtools-extension",
     "description":"Chrome/Firefox DevTools extension for performance profiling React and Vue apps. Captures render timings, prop diffs, component tree snapshots, and generates actionable reports.",
     "language":"TypeScript","stars":"2.8k",
     "score":78,"grade":"B","working":85,"recommendation":"Interview Further","complexity":77,"security":76,"docs":80,"maintainability":79,
     "tech_stack":["TypeScript","Chrome Extensions API","React","WebWorkers","IndexedDB"],
     "capabilities":["Browser APIs","Performance Profiling","Extension Development"],
     "red_flags":["Firefox compat layer incomplete","No automated UI tests"],"summary":"Useful developer tool with good UX; cross-browser support needs work."},
    {"name":"Federated Learning Framework","user":"gh_user_2","github":"federated-learning-framework",
     "description":"Privacy-preserving federated learning framework with differential privacy, secure aggregation, and support for heterogeneous client devices. Implements FedAvg and FedProx.",
     "language":"Python","stars":"3.6k",
     "score":86,"grade":"B","working":88,"recommendation":"Hire","complexity":94,"security":90,"docs":85,"maintainability":83,
     "tech_stack":["Python","PyTorch","gRPC","PySyft","Flower","NumPy"],
     "capabilities":["Federated Learning","Privacy Engineering","Distributed ML"],
     "red_flags":["Communication overhead not profiled"],"summary":"Cutting-edge privacy-preserving ML with solid theoretical grounding."},
    {"name":"Code Review Automation","user":"gh_user_3","github":"code-review-automation",
     "description":"AI-powered code review bot integrated with GitHub/GitLab. Uses static analysis + LLM to detect bugs, security issues, style violations, and generate fix suggestions as PR comments.",
     "language":"Python","stars":"6.2k",
     "score":87,"grade":"A","working":90,"recommendation":"Hire","complexity":85,"security":88,"docs":86,"maintainability":87,
     "tech_stack":["Python","GPT-4","Tree-sitter","GitHub Actions","FastAPI","PostgreSQL"],
     "capabilities":["LLM Integration","Static Analysis","DevOps Automation","AST Parsing"],
     "red_flags":["Rate limiting on GitHub API not fully handled"],"summary":"Practical AI tooling with good accuracy and clean GitHub integration."},
    {"name":"Database Query Planner","user":"gh_user_4","github":"query-planner",
     "description":"Custom SQL query planner and optimiser with cost-based optimisation, index selection, join reordering, and materialized view recommendations. Compatible with PostgreSQL protocol.",
     "language":"Rust","stars":"4.1k",
     "score":89,"grade":"A","working":91,"recommendation":"Strong Hire","complexity":98,"security":84,"docs":87,"maintainability":84,
     "tech_stack":["Rust","PostgreSQL Protocol","SQL Parser","LLVM"],
     "capabilities":["Database Internals","Query Optimisation","Rust","Compiler Design"],
     "red_flags":["Limited test coverage for edge-case queries"],"summary":"Deep database internals expertise with impressive query optimisation implementation."},
    {"name":"Observability Platform","user":"gh_user_5","github":"observability-platform",
     "description":"Unified observability platform combining logs, metrics, and traces with automatic correlation. Features anomaly detection, SLO monitoring, and intelligent alerting with Slack/PagerDuty.",
     "language":"Go","stars":"7.8k",
     "score":90,"grade":"A","working":94,"recommendation":"Strong Hire","complexity":91,"security":89,"docs":91,"maintainability":90,
     "tech_stack":["Go","Prometheus","Jaeger","Loki","Grafana","Kafka","ClickHouse"],
     "capabilities":["Observability","SRE","Distributed Tracing","Anomaly Detection"],
     "red_flags":[],"summary":"Production-ready observability stack with strong SRE practices."},
]

SAMPLE_RESUME_NOTE = """[SAMPLE PROJECT — Auto-generated resume placeholder]

This project was imported from GitHub as a sample portfolio entry.

Name: GitHub Developer
Title: Software Engineer
Skills: See project tech stack for relevant technologies.
Experience: Based on project complexity and code quality.

Note: This is a demonstration entry. Real candidate data would be
provided when a developer submits their own project and resume
through the DevTrustProtocol upload portal.

Contact: Extracted from actual resume upon real submission.
"""

# ──────────────────────────────────────────────────────────
# 6. ACCEPTED FILE TYPES
# ──────────────────────────────────────────────────────────
IMAGE_EXTENSIONS = {".png",".jpg",".jpeg",".gif",".webp",".bmp",".svg",".ico",".tiff"}
VIDEO_EXTENSIONS = {".mp4",".mov",".avi",".mkv",".webm",".wmv",".flv",".m4v"}
ACCEPTED_PROJECT_TYPES = [
    "zip","py","js","ts","jsx","tsx","html","css","json","yaml","yml",
    "toml","md","txt","csv","pdf","docx","xlsx","xls",
    "png","jpg","jpeg","gif","webp","bmp","svg",
    "mp4","mov","avi","mkv","webm","go","rs","java","c","cpp","h","rb","php","swift","kt",
]

# ──────────────────────────────────────────────────────────
# 7. DATA SANITIZER
# ──────────────────────────────────────────────────────────
class DataSanitizer:
    _PATTERNS = [
        (re.compile(r'(?i)(api[_\-]?key|apikey|access[_\-]?key)\s*[=:]\s*["\']?[\w\-]{16,}["\']?'), "API_KEY"),
        (re.compile(r'(?i)(secret|password|passwd|pwd|token|auth)\s*[=:]\s*["\']?[\w\-\@\#\$\%\!\&]{8,}["\']?'), "SECRET"),
        (re.compile(r'sk-[A-Za-z0-9]{32,}'), "OPENAI_KEY"),
        (re.compile(r'(?i)bearer\s+[A-Za-z0-9\-_\.]{20,}'), "BEARER_TOKEN"),
        (re.compile(r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'), "EMAIL"),
        (re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b'), "IP_ADDR"),
        (re.compile(r'(?i)(mongodb|postgres|mysql|redis|amqp)://[^\s\'"]+'), "DB_URL"),
        (re.compile(r'(?i)project[_\-]?id\s*[=:]\s*["\']?[\w\-]{6,}["\']?'), "PROJECT_ID"),
    ]
    _PROP = re.compile(r'#\s*(proprietary|confidential|internal|do not share)', re.I)

    def __init__(self):
        self._redacted = 0
        self._flagged  = 0

    def _token(self, label, value):
        h = hashlib.md5(value.encode()).hexdigest()[:6].upper()
        return f"[REDACTED_{label}_{h}]"

    def sanitize_text(self, text):
        if not isinstance(text, str): return text
        result = text
        for pattern, label in self._PATTERNS:
            def repl(m, lbl=label, pat=pattern):
                self._redacted += 1
                return self._token(lbl, m.group(0))
            result = pattern.sub(repl, result)
        lines = []
        for line in result.splitlines():
            if self._PROP.search(line):
                self._flagged += 1
            else:
                lines.append(line)
        return "\n".join(lines)

    def sanitize_resume_for_download(self, text):
        result = text
        for pattern, label in self._PATTERNS:
            if label not in ("EMAIL","IP_ADDR"):
                def repl(m, lbl=label):
                    return self._token(lbl, m.group(0))
                result = pattern.sub(repl, result)
        return result

    def get_report(self):
        return {"redacted_values": self._redacted, "flagged_lines": self._flagged}

# ──────────────────────────────────────────────────────────
# 8. PROJECT LOADER
# ──────────────────────────────────────────────────────────
class ProjectLoader:
    def load_many(self, uploaded_files):
        if not uploaded_files:
            return None
        if len(uploaded_files) == 1 and uploaded_files[0].name.endswith(".zip"):
            return self._load_zip(uploaded_files[0])
        return self._load_flat(uploaded_files)

    def _read_file(self, name, data):
        ext = pathlib.Path(name).suffix.lower()
        if ext in IMAGE_EXTENSIONS:
            b64 = base64.b64encode(data).decode()
            mime = f"image/{ext.lstrip('.')}" if ext != ".svg" else "image/svg+xml"
            return f"[IMAGE:{mime}:{b64}]"
        if ext in VIDEO_EXTENSIONS:
            return f"[VIDEO:{ext}:{len(data)//1024}KB]"
        if ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(stream=data, filetype="pdf")
                return "\n".join(p.get_text() for p in doc)
            except Exception: return "[PDF: could not extract text]"
        if ext == ".docx":
            try:
                import docx
                doc = docx.Document(io.BytesIO(data))
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception: return "[DOCX: could not extract text]"
        if ext in (".xlsx",".xls"):
            try:
                import openpyxl
                wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True)
                rows = []
                for ws in wb.worksheets:
                    rows.append(f"[Sheet: {ws.title}]")
                    for row in ws.iter_rows(values_only=True):
                        rows.append("\t".join(str(c) if c else "" for c in row))
                return "\n".join(rows)
            except Exception: return "[XLSX: could not extract]"
        try:
            return data.decode("utf-8", errors="replace")
        except Exception:
            return "[Binary file]"

    def _load_zip(self, f):
        f.seek(0)
        files, images, videos = {}, {}, {}
        try:
            with zipfile.ZipFile(io.BytesIO(f.read())) as zf:
                for name in zf.namelist():
                    if name.endswith("/"): continue
                    data = zf.read(name)
                    txt = self._read_file(name, data)
                    files[name] = txt
                    ext = pathlib.Path(name).suffix.lower()
                    if ext in IMAGE_EXTENSIONS: images[name] = txt
                    if ext in VIDEO_EXTENSIONS: videos[name] = {"raw": data}
        except Exception as e:
            return None
        return {"source": f.name, "type": "zip", "files": files,
                "images": images, "videos": videos,
                "file_count": len(files), "uploaded_count": 1,
                "summary": f"ZIP: {len(files)} files extracted"}

    def _load_flat(self, uploaded_files):
        files, images, videos = {}, {}, {}
        for uf in uploaded_files:
            uf.seek(0)
            data = uf.read()
            txt = self._read_file(uf.name, data)
            files[uf.name] = txt
            ext = pathlib.Path(uf.name).suffix.lower()
            if ext in IMAGE_EXTENSIONS: images[uf.name] = txt
            if ext in VIDEO_EXTENSIONS: videos[uf.name] = {"raw": data}
        return {"source": ", ".join(f.name for f in uploaded_files), "type": "files",
                "files": files, "images": images, "videos": videos,
                "file_count": len(files), "uploaded_count": len(uploaded_files),
                "summary": f"{len(files)} file(s) loaded"}

def extract_resume_text(uploaded_file):
    if not uploaded_file: return ""
    ext = pathlib.Path(uploaded_file.name).suffix.lower()
    if ext == ".pdf":
        try:
            import fitz
            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            return "\n".join(p.get_text() for p in doc)
        except Exception as e: st.error(f"PDF error: {e}"); return ""
    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e: st.error(f"DOCX error: {e}"); return ""
    if ext == ".json":
        try:
            data = json.load(uploaded_file)
            lines = []
            def flatten(obj, prefix=""):
                if isinstance(obj,dict):
                    for k,v in obj.items(): flatten(v, f"{prefix}{k}: " if not prefix else f"{prefix} > {k}: ")
                elif isinstance(obj,list):
                    for item in obj: flatten(item, prefix)
                else:
                    if str(obj).strip(): lines.append(f"{prefix}{obj}")
            flatten(data)
            return "\n".join(lines)
        except Exception as e: st.error(f"JSON error: {e}"); return ""
    return ""

# ──────────────────────────────────────────────────────────
# 9. RESUME PARSER
# ──────────────────────────────────────────────────────────
class ResumeParser:
    SYSTEM = """You are a resume parser. Return ONLY valid JSON with these keys:
{"name","email","phone","location","title","years_experience","top_skills","education","linkedin","github","summary_bio"}
Return ONLY JSON, no markdown, no explanation."""

    def __init__(self, provider, api_key, model):
        self.provider=provider; self.api_key=api_key; self.model=model

    def parse(self, text):
        prompt = f"Parse this resume:\n\n{text[:4000]}"
        raw = self._call(prompt)
        try:
            clean = raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
            return json.loads(clean)
        except Exception:
            return {"name":"Unknown","email":"","phone":"","location":"","title":"Developer",
                    "years_experience":0,"top_skills":[],"education":"N/A",
                    "linkedin":"Not found","github":"Not found","summary_bio":""}

    def _call(self, prompt):
        if self.provider == "Groq":
            from groq import Groq
            client = Groq(api_key=self.api_key)
            resp = client.chat.completions.create(
                model=self.model, messages=[{"role":"system","content":self.SYSTEM},
                                             {"role":"user","content":prompt}], max_tokens=800)
            return resp.choices[0].message.content
        else:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=self.api_key)
            resp = client.models.generate_content(
                model=self.model, contents=[prompt],
                config=types.GenerateContentConfig(system_instruction=self.SYSTEM, max_output_tokens=800))
            return resp.text

# ──────────────────────────────────────────────────────────
# 10. TRUST SCORER
# ──────────────────────────────────────────────────────────
class TrustScorer:
    """
    Adaptive scoring engine — adjusts evaluation dimensions & weights
    based on the declared project type.

    Project Types → Scoring Profile:
    ─────────────────────────────────────────────────────────────────
    Software Development    → original 9-dim engineering rubric
    Research & Analytics    → methodology, insight quality, reproducibility, visualisation
    Data Science / ML       → model quality, feature engineering, evaluation rigour
    Design & UX             → user research depth, prototype fidelity, accessibility
    Business / Strategy     → problem framing, evidence quality, actionability
    Other / General         → balanced across all areas
    ─────────────────────────────────────────────────────────────────
    """

    # ── Per-type system prompts ────────────────────────────────────
    _PROMPTS = {}

    _PROMPTS["Software Development"] = """You are a senior engineering lead performing a rigorous technical review.

PROJECT TYPE: Software Development / Engineering

Evaluate across these 9 dimensions (each 0-100):

1. functional_correctness (22%) — Does it run? Inputs validated? Edge cases handled?
2. code_quality (18%) — Naming, DRY, function length, style consistency, clean structure
3. security_practices (15%) — No hardcoded secrets, safe inputs, correct auth patterns
4. architecture_design (12%) — Separation of concerns, SOLID, modularity, scalability
5. test_reliability (10%) — Unit/integration tests, CI configured, critical path coverage
6. documentation (8%) — README, inline comments, API docs, setup instructions
7. dependency_hygiene (6%) — Pinned versions, lock file, no abandoned packages
8. performance_awareness (5%) — No O(n²) on large data, appropriate caching, async I/O
9. code_originality (4%) — Genuine problem-solving, not pure boilerplate or copy-paste

trust_score = round(fc*0.22 + cq*0.18 + sp*0.15 + ad*0.12 + tr*0.10 + doc*0.08 + dh*0.06 + pa*0.05 + co*0.04)
Strong Hire: score>=82 AND functional_correctness>=75 AND security_practices>=70
Hire: score>=68 AND functional_correctness>=60
Interview Further: score>=48 OR (score<68 AND functional_correctness>=50)
Decline: everything else"""

    _PROMPTS["Research & Analytics"] = """You are a senior research director evaluating an analytics or research project.

PROJECT TYPE: Research & Analytics

Evaluate across these 9 dimensions (each 0-100):

1. research_rigor (25%) — Methodology soundness, hypothesis clarity, bias controls, reproducibility
2. insight_quality (20%) — Depth of findings, novelty, accuracy of conclusions from data
3. data_handling (15%) — Data quality checks, missing value handling, source credibility, validation
4. visualisation (12%) — Clarity of charts, appropriate chart types, accessibility, narrative flow
5. documentation (10%) — Report structure, methodology notes, assumptions stated, references
6. statistical_validity (8%) — Correct statistical tests, significance interpretation, effect sizes
7. reproducibility (5%) — Code/queries to reproduce results, version-controlled data pipeline
8. code_quality (3%) — Clean analysis scripts (if any), readable notebooks, clear variable names
9. originality (2%) — Novel framing, unique dataset use, non-obvious insights

trust_score = round(rr*0.25 + iq*0.20 + dh*0.15 + vis*0.12 + doc*0.10 + sv*0.08 + rep*0.05 + cq*0.03 + orig*0.02)
Strong Hire: score>=82 AND research_rigor>=75 AND insight_quality>=70
Hire: score>=68 AND research_rigor>=60
Interview Further: score>=48 OR (score<68 AND research_rigor>=50)
Decline: everything else

Map your dimension scores to these JSON keys:
functional_correctness=research_rigor, code_quality=code_quality, security_practices=statistical_validity,
architecture_design=data_handling, test_reliability=reproducibility, documentation=documentation,
dependency_hygiene=originality, performance_awareness=visualisation, code_originality=insight_quality"""

    _PROMPTS["Data Science / ML"] = """You are a senior ML engineer and data scientist evaluating a machine learning project.

PROJECT TYPE: Data Science / Machine Learning

Evaluate across these 9 dimensions (each 0-100):

1. model_quality (22%) — Algorithm choice, hyperparameter tuning, train/val/test split correctness
2. feature_engineering (18%) — Feature selection quality, handling of categorical/missing data, leakage prevention
3. evaluation_rigour (15%) — Appropriate metrics, cross-validation, baseline comparisons, overfitting analysis
4. data_pipeline (12%) — Reproducible preprocessing, data versioning, clean ETL
5. code_quality (10%) — Clean notebooks/scripts, DRY, modular code, readable
6. documentation (8%) — Problem statement, model card, results interpretation, limitations noted
7. reproducibility (6%) — Requirements pinned, seeds set, steps to reproduce results documented
8. visualisation (5%) — Confusion matrices, learning curves, feature importance charts
9. originality (4%) — Novel approach, creative problem framing, non-standard solutions

trust_score = round(mq*0.22 + fe*0.18 + er*0.15 + dp*0.12 + cq*0.10 + doc*0.08 + rep*0.06 + vis*0.05 + orig*0.04)
Strong Hire: score>=82 AND model_quality>=75 AND evaluation_rigour>=70
Hire: score>=68 AND model_quality>=60
Interview Further: score>=48 OR (score<68 AND model_quality>=50)
Decline: everything else

Map to JSON keys:
functional_correctness=model_quality, code_quality=code_quality, security_practices=evaluation_rigour,
architecture_design=data_pipeline, test_reliability=reproducibility, documentation=documentation,
dependency_hygiene=originality, performance_awareness=visualisation, code_originality=feature_engineering"""

    _PROMPTS["Design & UX"] = """You are a senior UX director and design lead evaluating a design or UX project.

PROJECT TYPE: Design & UX

Evaluate across these 9 dimensions (each 0-100):

1. user_research_depth (22%) — Research methods used, sample sizes, insight synthesis quality
2. usability (20%) — Clarity of flows, cognitive load, error prevention, onboarding ease
3. visual_design (15%) — Consistency, typography, colour accessibility (WCAG), polish
4. prototype_fidelity (12%) — Detail level, interaction completeness, realistic content
5. problem_framing (10%) — Clear problem statement, user personas, job-to-be-done clarity
6. documentation (8%) — Design rationale, component notes, accessibility decisions recorded
7. iteration_evidence (6%) — User testing conducted, changes made based on feedback
8. accessibility (5%) — Contrast ratios, keyboard nav, screen-reader consideration
9. originality (2%) — Creative solutions, non-obvious design patterns, fresh perspective

trust_score = round(urd*0.22 + usa*0.20 + vd*0.15 + pf*0.12 + prob*0.10 + doc*0.08 + ie*0.06 + acc*0.05 + orig*0.02)
Strong Hire: score>=82 AND user_research_depth>=75 AND usability>=70
Hire: score>=68 AND user_research_depth>=60
Interview Further: score>=48 OR (score<68 AND user_research_depth>=50)
Decline: everything else

Map to JSON keys:
functional_correctness=user_research_depth, code_quality=usability, security_practices=accessibility,
architecture_design=problem_framing, test_reliability=iteration_evidence, documentation=documentation,
dependency_hygiene=originality, performance_awareness=visual_design, code_originality=prototype_fidelity"""

    _PROMPTS["Business / Strategy"] = """You are a senior strategy consultant evaluating a business or strategy project.

PROJECT TYPE: Business / Strategy

Evaluate across these 9 dimensions (each 0-100):

1. problem_framing (22%) — Is the business problem clearly defined? Scope appropriate? Root cause identified?
2. evidence_quality (20%) — Are claims backed by data? Sources credible? Quantified where possible?
3. insight_depth (15%) — Are insights non-obvious? Does analysis go beyond surface-level observations?
4. actionability (12%) — Are recommendations specific, feasible, prioritised, and measurable?
5. structure (10%) — Logical flow, MECE principles, clear executive summary, stakeholder-appropriate
6. documentation (8%) — Sources cited, assumptions listed, methodology explained
7. market_awareness (6%) — Competitive landscape considered, macro factors addressed, timing sensible
8. risk_analysis (5%) — Risks identified, mitigations proposed, sensitivity analysis if applicable
9. originality (2%) — Novel framing, creative strategic options, non-boilerplate recommendations

trust_score = round(pf*0.22 + eq*0.20 + id*0.15 + act*0.12 + str*0.10 + doc*0.08 + ma*0.06 + ra*0.05 + orig*0.02)
Strong Hire: score>=82 AND problem_framing>=75 AND evidence_quality>=70
Hire: score>=68 AND problem_framing>=60
Interview Further: score>=48 OR (score<68 AND problem_framing>=50)
Decline: everything else

Map to JSON keys:
functional_correctness=problem_framing, code_quality=structure, security_practices=risk_analysis,
architecture_design=market_awareness, test_reliability=actionability, documentation=documentation,
dependency_hygiene=originality, performance_awareness=evidence_quality, code_originality=insight_depth"""

    _PROMPTS["Other / General"] = """You are an expert evaluator performing a holistic project assessment.

PROJECT TYPE: General / Other

Evaluate across these 9 dimensions (each 0-100):

1. functional_correctness (22%) — Does the project achieve its stated goals? Is the output correct?
2. quality (18%) — Overall quality of execution — clarity, completeness, correctness
3. methodology (15%) — Is the approach sound? Are best practices for this domain followed?
4. structure (12%) — Organisation, logical flow, modularity or layering where applicable
5. reliability (10%) — Can the work be trusted? Is it consistent and validated?
6. documentation (8%) — Is the work explained clearly? Can someone else understand or reproduce it?
7. hygiene (6%) — Clean artefacts, no unnecessary clutter, dependencies/sources cited
8. performance (5%) — Efficient use of resources or time; no obvious waste
9. originality (4%) — Genuine effort, not generic or templated

trust_score = round(fc*0.22 + q*0.18 + meth*0.15 + str*0.12 + rel*0.10 + doc*0.08 + hyg*0.06 + perf*0.05 + orig*0.04)
Strong Hire: score>=82 AND functional_correctness>=75
Hire: score>=68 AND functional_correctness>=60
Interview Further: score>=48 OR (score<68 AND functional_correctness>=50)
Decline: everything else

Map to JSON keys:
functional_correctness=functional_correctness, code_quality=quality, security_practices=methodology,
architecture_design=structure, test_reliability=reliability, documentation=documentation,
dependency_hygiene=hygiene, performance_awareness=performance, code_originality=originality"""

    # ── Shared JSON output spec appended to every prompt ──────────
    _JSON_SPEC = """

Return ONLY valid JSON — no markdown, no preamble:
{
  "trust_score": <0-100 weighted composite>,
  "grade": <"A"|"B"|"C"|"D"|"F">,
  "working_percentage": <0-100 — what % of stated goals is implemented/working>,
  "recommendation": <"Strong Hire"|"Hire"|"Interview Further"|"Decline">,
  "functional_correctness": <0-100>,
  "code_quality": <0-100>,
  "security_practices": <0-100>,
  "architecture_design": <0-100>,
  "test_reliability": <0-100>,
  "documentation": <0-100>,
  "dependency_hygiene": <0-100>,
  "performance_awareness": <0-100>,
  "code_originality": <0-100>,
  "tech_stack_summary": [<tools, languages, frameworks, or methods detected>],
  "key_capabilities": [<3-6 specific skills demonstrated>],
  "red_flags": [<concrete actionable concerns, empty if none>],
  "strengths": [<3-5 genuine strengths observed>],
  "executive_summary": "<3 sentences: what the project does, quality level, hiring signal>"
}"""

    # ── Dimension weight maps per type ─────────────────────────────
    _WEIGHTS_BY_TYPE = {
        "Software Development": {
            "functional_correctness":0.22,"code_quality":0.18,"security_practices":0.15,
            "architecture_design":0.12,"test_reliability":0.10,"documentation":0.08,
            "dependency_hygiene":0.06,"performance_awareness":0.05,"code_originality":0.04,
        },
        "Research & Analytics": {
            "functional_correctness":0.25,"code_originality":0.20,"security_practices":0.08,
            "architecture_design":0.15,"test_reliability":0.05,"documentation":0.10,
            "dependency_hygiene":0.02,"performance_awareness":0.12,"code_quality":0.03,
        },
        "Data Science / ML": {
            "functional_correctness":0.22,"code_originality":0.18,"security_practices":0.15,
            "architecture_design":0.12,"test_reliability":0.06,"documentation":0.08,
            "dependency_hygiene":0.04,"performance_awareness":0.05,"code_quality":0.10,
        },
        "Design & UX": {
            "functional_correctness":0.22,"code_quality":0.20,"security_practices":0.05,
            "architecture_design":0.10,"test_reliability":0.06,"documentation":0.08,
            "dependency_hygiene":0.02,"performance_awareness":0.15,"code_originality":0.12,
        },
        "Business / Strategy": {
            "functional_correctness":0.22,"code_quality":0.10,"security_practices":0.05,
            "architecture_design":0.06,"test_reliability":0.12,"documentation":0.08,
            "dependency_hygiene":0.02,"performance_awareness":0.20,"code_originality":0.15,
        },
    }
    _WEIGHTS_DEFAULT = {
        "functional_correctness":0.22,"code_quality":0.18,"security_practices":0.12,
        "architecture_design":0.12,"test_reliability":0.10,"documentation":0.08,
        "dependency_hygiene":0.06,"performance_awareness":0.08,"code_originality":0.04,
    }

    # ── Dimension labels per type (for UI display) ─────────────────
    DIMENSION_LABELS = {
        "Software Development": [
            ("functional_correctness","Functional Correctness","✅",22,"Does it run? Inputs validated?","#16a34a"),
            ("code_quality","Code Quality","🖊️",18,"Clean, readable, maintainable","#2563eb"),
            ("security_practices","Security Practices","🔒",15,"No secrets, safe inputs, auth","#dc2626"),
            ("architecture_design","Architecture & Design","🏗️",12,"SOLID, modularity, scalability","#7c3aed"),
            ("test_reliability","Test Reliability","🧪",10,"Tests, CI, critical coverage","#0891b2"),
            ("documentation","Documentation","📖",8,"README, comments, API docs","#6b7280"),
            ("dependency_hygiene","Dependency Hygiene","📦",6,"Pinned versions, no abandoned pkgs","#d97706"),
            ("performance_awareness","Performance","⚡",5,"No O(n²), caching, async I/O","#ea580c"),
            ("code_originality","Code Originality","✨",4,"Genuine problem-solving","#9333ea"),
        ],
        "Research & Analytics": [
            ("functional_correctness","Research Rigor","🔬",25,"Methodology, hypothesis, bias controls","#16a34a"),
            ("code_originality","Insight Quality","💡",20,"Depth of findings, novelty, accuracy","#2563eb"),
            ("architecture_design","Data Handling","🗄️",15,"Quality checks, sources, validation","#7c3aed"),
            ("performance_awareness","Visualisation","📊",12,"Chart clarity, appropriate types, narrative","#ea580c"),
            ("documentation","Documentation","📖",10,"Report structure, assumptions, references","#6b7280"),
            ("security_practices","Statistical Validity","📐",8,"Correct tests, significance, effect sizes","#dc2626"),
            ("test_reliability","Reproducibility","♻️",5,"Code to reproduce, versioned pipeline","#0891b2"),
            ("code_quality","Code Quality","🖊️",3,"Clean scripts, readable notebooks","#d97706"),
            ("dependency_hygiene","Originality","✨",2,"Novel framing, unique dataset use","#9333ea"),
        ],
        "Data Science / ML": [
            ("functional_correctness","Model Quality","🤖",22,"Algorithm choice, tuning, split correctness","#16a34a"),
            ("code_originality","Feature Engineering","⚙️",18,"Feature selection, leakage prevention","#2563eb"),
            ("security_practices","Evaluation Rigor","📏",15,"Metrics, cross-val, baseline comparison","#dc2626"),
            ("architecture_design","Data Pipeline","🗄️",12,"Reproducible preprocessing, clean ETL","#7c3aed"),
            ("code_quality","Code Quality","🖊️",10,"Clean notebooks/scripts, modular","#0891b2"),
            ("documentation","Documentation","📖",8,"Model card, results interpretation","#6b7280"),
            ("test_reliability","Reproducibility","♻️",6,"Pinned seeds, steps to reproduce","#d97706"),
            ("performance_awareness","Visualisation","📊",5,"Learning curves, confusion matrices","#ea580c"),
            ("dependency_hygiene","Originality","✨",4,"Novel approach, creative framing","#9333ea"),
        ],
        "Design & UX": [
            ("functional_correctness","User Research","👥",22,"Methods, sample size, insight synthesis","#16a34a"),
            ("code_quality","Usability","🖱️",20,"Clarity, cognitive load, error prevention","#2563eb"),
            ("performance_awareness","Visual Design","🎨",15,"Consistency, typography, colour/WCAG","#ea580c"),
            ("code_originality","Prototype Fidelity","🖼️",12,"Detail level, interaction completeness","#9333ea"),
            ("architecture_design","Problem Framing","🎯",10,"Clear problem, personas, JTBD","#7c3aed"),
            ("documentation","Documentation","📖",8,"Design rationale, component notes","#6b7280"),
            ("test_reliability","Iteration Evidence","🔄",6,"User testing, changes from feedback","#0891b2"),
            ("security_practices","Accessibility","♿",5,"Contrast, keyboard nav, screen reader","#d97706"),
            ("dependency_hygiene","Originality","✨",2,"Creative solutions, fresh patterns","#9333ea"),
        ],
        "Business / Strategy": [
            ("functional_correctness","Problem Framing","🎯",22,"Clear problem, scope, root cause","#16a34a"),
            ("performance_awareness","Evidence Quality","📊",20,"Data-backed claims, credible sources","#ea580c"),
            ("code_originality","Insight Depth","💡",15,"Non-obvious insights, beyond surface","#2563eb"),
            ("test_reliability","Actionability","⚡",12,"Specific, feasible, prioritised recs","#0891b2"),
            ("code_quality","Structure","📋",10,"MECE, logical flow, exec summary","#d97706"),
            ("documentation","Documentation","📖",8,"Sources, assumptions, methodology","#6b7280"),
            ("architecture_design","Market Awareness","🌍",6,"Competitive landscape, macro factors","#7c3aed"),
            ("security_practices","Risk Analysis","⚠️",5,"Risks identified, mitigations proposed","#dc2626"),
            ("dependency_hygiene","Originality","✨",2,"Novel framing, creative options","#9333ea"),
        ],
    }
    _LABELS_DEFAULT = [
        ("functional_correctness","Correctness / Goals","✅",22,"Achieves stated goals correctly","#16a34a"),
        ("code_quality","Quality","⭐",18,"Overall execution quality","#2563eb"),
        ("security_practices","Methodology","🔬",15,"Sound approach, best practices","#dc2626"),
        ("architecture_design","Structure","🏗️",12,"Organisation and logical flow","#7c3aed"),
        ("test_reliability","Reliability","🛡️",10,"Consistency, validation, trustworthiness","#0891b2"),
        ("documentation","Documentation","📖",8,"Clarity, explanations, reproducibility","#6b7280"),
        ("dependency_hygiene","Hygiene","🧹",6,"Clean artefacts, no clutter","#d97706"),
        ("performance_awareness","Performance","⚡",5,"Efficient, minimal waste","#ea580c"),
        ("code_originality","Originality","✨",4,"Genuine effort, non-generic","#9333ea"),
    ]

    PROJECT_TYPES = [
        "Software Development",
        "Research & Analytics",
        "Data Science / ML",
        "Design & UX",
        "Business / Strategy",
        "Other / General",
    ]

    def __init__(self, provider, api_key, model):
        self.provider = provider
        self.api_key  = api_key
        self.model    = model

    def _get_weights(self, project_type: str) -> dict:
        return self._WEIGHTS_BY_TYPE.get(project_type, self._WEIGHTS_DEFAULT)

    def get_dimension_labels(self, project_type: str) -> list:
        return self.DIMENSION_LABELS.get(project_type, self._LABELS_DEFAULT)

    def _calc_trust_score(self, raw: dict, project_type: str) -> dict:
        weights = self._get_weights(project_type)
        composite = sum(raw.get(dim, 50) * w for dim, w in weights.items())
        trust = max(0, min(100, round(composite)))
        grade = "F"
        if trust >= 85: grade = "A"
        elif trust >= 70: grade = "B"
        elif trust >= 55: grade = "C"
        elif trust >= 40: grade = "D"
        fc  = raw.get("functional_correctness", 50)
        sec = raw.get("security_practices", 50)
        if trust >= 82 and fc >= 75 and sec >= 70:
            rec = "Strong Hire"
        elif trust >= 68 and fc >= 60:
            rec = "Hire"
        elif trust >= 48 or (trust < 68 and fc >= 50):
            rec = "Interview Further"
        else:
            rec = "Decline"
        raw.update({"trust_score": trust, "grade": grade, "recommendation": rec,
                    "project_type": project_type})
        # backward-compat aliases
        for alias, real in [("complexity","architecture_design"),("test_coverage","test_reliability"),
                            ("maintainability","code_quality"),("performance","performance_awareness"),
                            ("dependencies","dependency_hygiene"),("project_completeness","functional_correctness")]:
            if alias not in raw:
                raw[alias] = raw.get(real, 50)
        return raw

    def score(self, project_data: dict, profile: dict, project_type: str = "Software Development",
              description: str = "") -> dict:
        files = project_data.get("files", {})
        snippets, tree_lines = [], []
        for path, body in list(files.items())[:25]:
            tree_lines.append(path)
            if not body.startswith(("[IMAGE:", "[VIDEO:")):
                snippets.append(f"=== {path} ===\n{body[:2000]}")
        file_tree = "\n".join(sorted(tree_lines))
        candidate = (
            f"Candidate: {profile.get('name','?')}\n"
            f"Title: {profile.get('title','?')}\n"
            f"Skills: {', '.join(profile.get('top_skills',[])[:12])}\n"
            f"Experience: {profile.get('years_experience','?')} years"
        )
        desc_block = f"\nPROJECT DESCRIPTION (from author):\n{description}\n" if description.strip() else ""
        code_block = "\n\n".join(snippets[:12])
        system_prompt = self._PROMPTS.get(project_type, self._PROMPTS["Other / General"]) + self._JSON_SPEC
        prompt = (
            f"PROJECT: {project_data.get('source','?')}\n"
            f"PROJECT TYPE: {project_type}\n"
            f"FILE COUNT: {len(files)}\n"
            f"{desc_block}\n"
            f"FILE TREE:\n{file_tree}\n\n"
            f"{candidate}\n\n"
            f"FILE CONTENTS:\n{code_block}"
        )
        raw = self._call(prompt, system_prompt)
        try:
            clean = re.sub(r"^```(?:json)?\n?","",raw.strip())
            clean = re.sub(r"\n?```$","",clean.rstrip())
            parsed = json.loads(clean.strip())
        except Exception:
            parsed = {
                "trust_score":50,"grade":"C","working_percentage":50,
                "recommendation":"Interview Further","functional_correctness":50,
                "code_quality":50,"security_practices":50,"architecture_design":50,
                "test_reliability":50,"documentation":50,"dependency_hygiene":50,
                "performance_awareness":50,"code_originality":50,
                "tech_stack_summary":[],"key_capabilities":[],"strengths":[],
                "red_flags":["Could not parse AI response"],
                "executive_summary":"Automated scoring returned a parse error.",
            }
        return self._calc_trust_score(parsed, project_type)

    def _call(self, prompt: str, system_prompt: str) -> str:
        if self.provider == "Groq":
            from groq import Groq
            client = Groq(api_key=self.api_key)
            resp = client.chat.completions.create(
                model=self.model,
                messages=[{"role":"system","content":system_prompt},
                           {"role":"user","content":prompt}],
                max_tokens=1400, temperature=0.1,
            )
            return resp.choices[0].message.content
        else:
            from google import genai
            from google.genai import types
            client = genai.Client(api_key=self.api_key)
            resp = client.models.generate_content(
                model=self.model, contents=[prompt],
                config=types.GenerateContentConfig(
                    system_instruction=system_prompt, max_output_tokens=1400, temperature=0.1),
            )
            return resp.text


# ──────────────────────────────────────────────────────────
# 11. RESUME DOWNLOAD
# ──────────────────────────────────────────────────────────
def safe_resume_download(resume_text, profile, is_recruiter,
                         resume_bytes=None, resume_ext=".txt", resume_filename="resume"):
    try:
        if not resume_text:
            return
        ext_to_mime = {
            ".pdf":  "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".json": "application/json",
            ".txt":  "text/plain",
        }
        if is_recruiter:
            if resume_bytes and resume_ext in (".pdf",".docx"):
                st.download_button("⬇️ Download Resume", data=resume_bytes,
                    file_name=f"candidate_{resume_filename}{resume_ext}",
                    mime=ext_to_mime.get(resume_ext,"application/octet-stream"),
                    use_container_width=True)
            else:
                san = DataSanitizer()
                clean = san.sanitize_resume_for_download(resume_text)
                st.download_button("⬇️ Download Resume", data=clean.encode(),
                    file_name=f"candidate_{resume_filename}_sanitized.txt",
                    mime="text/plain", use_container_width=True)
        else:
            if resume_bytes:
                st.download_button("⬇️ Download Resume", data=resume_bytes,
                    file_name=f"{resume_filename}{resume_ext}",
                    mime=ext_to_mime.get(resume_ext,"application/octet-stream"),
                    use_container_width=True)
            else:
                st.download_button("⬇️ Download Resume", data=resume_text.encode(),
                    file_name=f"{resume_filename}.txt", mime="text/plain",
                    use_container_width=True)
    except Exception:
        pass

# ──────────────────────────────────────────────────────────
# 12. REPORT GENERATOR
# ──────────────────────────────────────────────────────────
class ReportGenerator:
    def render_candidate_profile(self, profile, resume_text, role="User"):
        name     = profile.get("name","Unknown Candidate")
        title    = profile.get("title","")
        email    = profile.get("email","")
        phone    = profile.get("phone","")
        loc      = profile.get("location","")
        skills   = profile.get("top_skills",[])
        edu      = profile.get("education","")
        bio      = profile.get("summary_bio","")
        exp      = profile.get("years_experience","")
        linkedin = profile.get("linkedin","")
        github   = profile.get("github","")

        # ── Profile header card ──
        meta_parts = []
        if title: meta_parts.append(title)
        if loc:   meta_parts.append(f"📍 {loc}")
        if exp:   meta_parts.append(f"⏱ {exp} yrs exp")
        meta_str = "  ·  ".join(meta_parts)

        contact_parts = []
        if email:    contact_parts.append(f"<span style='margin-right:16px;'>📧 {email}</span>")
        if phone:    contact_parts.append(f"<span style='margin-right:16px;'>📱 {phone}</span>")
        if linkedin: contact_parts.append(f"<a href='{linkedin}' target='_blank' style='color:#2563eb;margin-right:16px;'>🔗 LinkedIn</a>")
        if github:   contact_parts.append(f"<a href='{github}' target='_blank' style='color:#2563eb;margin-right:16px;'>🐙 GitHub</a>")
        contact_html = "".join(contact_parts)

        skill_chips = "".join(
            f"<span style='background:#eff6ff;color:#2563eb;border:1px solid #bfdbfe;"
            f"padding:3px 10px;border-radius:5px;font-size:12px;font-weight:500;"
            f"display:inline-block;margin:3px 4px 3px 0;'>{s}</span>"
            for s in skills
        )

        st.markdown(
            f"<div style='background:#ffffff;border:1px solid #e2e8f0;border-radius:10px;"
            f"padding:20px 22px;box-shadow:0 1px 3px rgba(0,0,0,.07);margin-bottom:16px;'>"
            # Name + title row
            f"<div style='display:flex;align-items:flex-start;justify-content:space-between;flex-wrap:wrap;gap:8px;margin-bottom:12px;'>"
            f"<div>"
            f"<div style='font-size:20px;font-weight:700;color:#111827;margin-bottom:3px;'>{name}</div>"
            + (f"<div style='font-size:13px;color:#6b7280;'>{meta_str}</div>" if meta_str else "")
            + f"</div></div>"
            # Contact row
            + (f"<div style='font-size:12px;color:#374151;margin-bottom:12px;"
               f"padding-bottom:12px;border-bottom:1px solid #e2e8f0;'>{contact_html}</div>"
               if contact_html else "")
            # Bio
            + (f"<div style='font-size:13px;color:#374151;line-height:1.7;"
               f"background:#f8fafc;border-radius:6px;padding:10px 12px;margin-bottom:12px;'>{bio}</div>"
               if bio else "")
            # Skills
            + (f"<div style='margin-bottom:10px;'>"
               f"<div style='font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.08em;"
               f"text-transform:uppercase;margin-bottom:6px;'>Skills</div>"
               f"<div>{skill_chips}</div></div>"
               if skills else "")
            # Education
            + (f"<div style='font-size:12px;color:#6b7280;margin-top:4px;'>🎓 {edu}</div>"
               if edu else "")
            + "</div>",
            unsafe_allow_html=True)

        # Resume download
        resume_bytes = st.session_state.get("resume_bytes")
        resume_ext   = st.session_state.get("resume_ext",".txt")
        resume_fname = st.session_state.get("resume_filename","resume")
        safe_resume_download(resume_text, profile,
            is_recruiter=(role=="Recruiter"),
            resume_bytes=resume_bytes, resume_ext=resume_ext, resume_filename=resume_fname)

    def render_trust_metrics(self, result):
        project_type = result.get("project_type", "Software Development")
        scorer       = TrustScorer.__new__(TrustScorer)  # no-init instance for label lookup
        dims         = scorer.get_dimension_labels(project_type)

        # ── Hero KPIs ────────────────────────────────────────────
        score   = result.get("trust_score", 0)
        grade   = result.get("grade", "?")
        working = result.get("working_percentage", 0)
        rec     = result.get("recommendation", "N/A")
        gc      = {"A":"#16a34a","B":"#2563eb","C":"#d97706","D":"#ea580c","F":"#dc2626"}.get(grade,"#6b7280")
        rc      = {"Strong Hire":"#16a34a","Hire":"#2563eb","Interview Further":"#d97706","Decline":"#dc2626"}.get(rec,"#6b7280")

        # Project type badge
        type_colors = {
            "Software Development": ("#2563eb","#eff6ff"),
            "Research & Analytics": ("#7c3aed","#f5f3ff"),
            "Data Science / ML":    ("#0891b2","#ecfeff"),
            "Design & UX":          ("#db2777","#fdf2f8"),
            "Business / Strategy":  ("#d97706","#fffbeb"),
            "Other / General":      ("#6b7280","#f8fafc"),
        }
        tc, tbg = type_colors.get(project_type, ("#6b7280","#f8fafc"))
        st.markdown(
            f"<div style='display:inline-flex;align-items:center;gap:6px;"
            f"background:{tbg};border:1px solid {tc}40;border-radius:20px;"
            f"padding:4px 14px;margin-bottom:16px;'>"
            f"<span style='font-size:11px;font-weight:700;color:{tc};letter-spacing:.06em;"
            f"text-transform:uppercase;'>{project_type}</span></div>",
            unsafe_allow_html=True)

        h1,h2,h3,h4 = st.columns(4)
        with h1: st.metric("🏆 Trust Score",   f"{score}/100")
        with h2: st.metric("📊 Grade",         grade)
        with h3: st.metric("⚙️ Working",        f"{working}%")
        with h4: st.metric("💼 Recommendation", rec)

        st.divider()

        # ── Dimension breakdown ─────────────────────────────────
        st.markdown(
            f"<p style='font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.09em;"
            f"text-transform:uppercase;margin-bottom:14px;'>Score Breakdown — {project_type}</p>",
            unsafe_allow_html=True)

        for key, label, ico, weight, tip, color in dims:
            val = result.get(key, 50)
            pct = max(0, min(100, int(val)))
            bar_color = "#16a34a" if pct >= 80 else "#d97706" if pct >= 55 else "#dc2626"
            st.markdown(
                f"<div style='margin-bottom:12px;'>"
                f"<div style='display:flex;justify-content:space-between;align-items:center;margin-bottom:5px;'>"
                f"<div style='display:flex;align-items:center;gap:8px;'>"
                f"<span style='font-size:14px;'>{ico}</span>"
                f"<div>"
                f"<span style='font-size:13px;font-weight:600;color:#111827;'>{label}</span>"
                f"<span style='font-size:10px;color:#9ca3af;margin-left:8px;'>{weight}% weight</span>"
                f"</div></div>"
                f"<div style='display:flex;align-items:center;gap:10px;'>"
                f"<span style='font-size:11px;color:#6b7280;'>{tip}</span>"
                f"<span style='font-size:16px;font-weight:700;color:{bar_color};"
                f"font-family:'JetBrains Mono',monospace;min-width:36px;text-align:right;'>{pct}</span>"
                f"</div></div>"
                f"<div style='background:#eef2f7;border-radius:6px;height:7px;overflow:hidden;'>"
                f"<div style='width:{pct}%;height:100%;background:{bar_color};border-radius:6px;'></div></div>"
                f"</div>",
                unsafe_allow_html=True)

        # ── Tech stack ──────────────────────────────────────────
        stack = result.get("tech_stack_summary", [])
        if stack:
            st.divider()
            st.markdown(
                "<p style='font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.09em;"
                "text-transform:uppercase;margin-bottom:10px;'>Tools & Technologies</p>",
                unsafe_allow_html=True)
            st.markdown(
                "<div style='display:flex;flex-wrap:wrap;gap:6px;'>"
                + "".join(
                    f"<span style='background:#eff6ff;color:#2563eb;border:1px solid #bfdbfe;"
                    f"padding:3px 10px;border-radius:5px;font-size:12px;font-weight:500;'>{t}</span>"
                    for t in stack)
                + "</div>",
                unsafe_allow_html=True)

        # ── Strengths ────────────────────────────────────────────
        strengths = result.get("strengths", [])
        if strengths:
            st.divider()
            st.markdown(
                "<p style='font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.09em;"
                "text-transform:uppercase;margin-bottom:10px;'>Strengths</p>",
                unsafe_allow_html=True)
            for s in strengths:
                st.markdown(
                    f"<div style='display:flex;gap:8px;align-items:flex-start;margin-bottom:6px;'>"
                    f"<span style='color:#16a34a;margin-top:1px;'>✓</span>"
                    f"<span style='font-size:13px;color:#374151;'>{s}</span></div>",
                    unsafe_allow_html=True)


    def render_summary_and_rec(self, result):
        rec   = result.get("recommendation","N/A")
        grade = result.get("grade","?")
        score = result.get("trust_score",0)
        st.markdown(f"### Assessment Summary")
        st.markdown(f"> {result.get('executive_summary','N/A')}")
        col1,col2,col3 = st.columns(3)
        with col1: col1.metric("Trust Score", f"{score}/100")
        with col2: col2.metric("Grade", grade)
        with col3: col3.metric("Recommendation", rec)
        caps = result.get("key_capabilities",[])
        if caps:
            st.markdown("**Key Capabilities:**")
            for c in caps: st.markdown(f"- {c}")
        flags = result.get("red_flags",[])
        if flags:
            st.markdown("**⚠️ Red Flags:**")
            for fl in flags: st.markdown(f"- {fl}")

# ══════════════════════════════════════════════════════════════
# ══════════════════════════════════════════════════════════════
# ██  UI LAYER — v5  (white, centered, sidebar always on, clean cards)
# ══════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════
# ██  UI LAYER v6 — persistent sidebar, login/logout/register
# ══════════════════════════════════════════════════════════════════════
import json as _json_mod, hashlib as _hash_mod

# ── Colour palette ─────────────────────────────────────────────────
C = {
    "bg":      "#f5f7fa",
    "s1":      "#ffffff",
    "s2":      "#f0f4f8",
    "s3":      "#e8edf3",
    "bd":      "#e2e8f0",
    "tx":      "#111827",
    "tx2":     "#374151",
    "sub":     "#6b7280",
    "blue":    "#2563eb",
    "bluelt":  "#eff6ff",
    "bluebdr": "#bfdbfe",
    "green":   "#16a34a",
    "greenlt": "#f0fdf4",
    "greenbdr":"#bbf7d0",
    "amber":   "#d97706",
    "amberlt": "#fffbeb",
    "amberbdr":"#fde68a",
    "red":     "#dc2626",
    "redlt":   "#fef2f2",
    "redbdr":  "#fecaca",
    "violet":  "#7c3aed",
    "violetlt":"#f5f3ff",
}

TYPE_COLOR = {
    "Software Development": "#2563eb",
    "Research & Analytics": "#7c3aed",
    "Data Science / ML":    "#0891b2",
    "Design & UX":          "#db2777",
    "Business / Strategy":  "#d97706",
    "Other / General":      "#6b7280",
}

# ── Dynamic user store (persists new registrations) ────────────────
_USERS_FILE = "devtrust_users.json"

def _load_dynamic_users():
    try:
        with open(_USERS_FILE) as f:
            return _json_mod.load(f)
    except Exception:
        return {}

def _save_dynamic_users(data):
    with open(_USERS_FILE, "w") as f:
        _json_mod.dump(data, f, indent=2)

def _hash_pw(pw):
    return _hash_mod.sha256(pw.encode()).hexdigest()

def _all_credentials():
    """Merge static CREDENTIALS with dynamic registered users."""
    dynamic = _load_dynamic_users()
    merged  = dict(CREDENTIALS)
    merged.update(dynamic)
    return merged

def register_user(username, password, display, role, email=""):
    """Register a new user. Returns (ok, message)."""
    if not username.strip():
        return False, "Username cannot be empty."
    if len(username) < 3:
        return False, "Username must be at least 3 characters."
    if not password or len(password) < 6:
        return False, "Password must be at least 6 characters."
    if username in _all_credentials():
        return False, "Username already exists. Please choose another."
    dynamic = _load_dynamic_users()
    dynamic[username] = {
        "password": password,   # store plain for consistency with static creds
        "role":     role,
        "display":  display or username,
        "email":    email,
        "registered": True,
    }
    _save_dynamic_users(dynamic)
    return True, "Account created successfully!"


# ── CSS injection ──────────────────────────────────────────────────
def _css():
    st.markdown("""<style>
/* ── Page background ── */
.stApp, [data-testid="stAppViewContainer"] { background-color: #f5f7fa; }
[data-testid="stMain"], [data-testid="stMainBlockContainer"],
.main, .block-container {
    background-color: #f5f7fa;
    padding: 0 !important;
    max-width: 100% !important;
}

/* ── SIDEBAR — always open, fixed, no collapse arrow ── */
section[data-testid="stSidebar"] {
    width: 260px !important;
    min-width: 260px !important;
    max-width: 260px !important;
    transform: translateX(0px) !important;
    background-color: #ffffff !important;
    border-right: 1px solid #e2e8f0 !important;
}
section[data-testid="stSidebar"] > div:first-child {
    width: 260px !important;
    background-color: #ffffff !important;
}
/* Hide the collapse/expand toggle arrow completely */
button[data-testid="baseButton-headerNoPadding"],
[data-testid="collapsedControl"],
.st-emotion-cache-1cypcdb,
button[kind="headerNoPadding"],
[data-testid="stSidebarCollapseButton"] {
    display: none !important;
}

/* ── Hide Streamlit chrome ── */
#MainMenu, footer, [data-testid="stHeader"],
[data-testid="stDecoration"], [data-testid="stToolbar"] { display: none !important; }

/* ── Inputs ── */
.stTextInput > div > div > input {
    border-radius: 7px !important;
    border: 1px solid #e2e8f0 !important;
    font-size: 14px !important;
}
.stTextInput > div > div > input:focus {
    border-color: #2563eb !important;
    box-shadow: 0 0 0 3px rgba(37,99,235,.12) !important;
}
.stTextArea > div > textarea {
    border-radius: 7px !important;
    border: 1px solid #e2e8f0 !important;
    font-size: 13px !important;
}

/* ── Metrics ── */
[data-testid="metric-container"] {
    background: #ffffff !important;
    border: 1px solid #e2e8f0 !important;
    border-radius: 10px !important;
    padding: 14px 16px !important;
    box-shadow: 0 1px 3px rgba(0,0,0,.07) !important;
}
[data-testid="stMetricLabel"] p {
    font-size: 11px !important;
    text-transform: uppercase !important;
    letter-spacing: .05em !important;
}
[data-testid="stMetricValue"] {
    font-size: 20px !important;
    font-weight: 700 !important;
}

/* ── File uploader ── */
[data-testid="stFileUploadDropzone"] {
    background: #f8fafc !important;
    border: 2px dashed #e2e8f0 !important;
    border-radius: 10px !important;
}

/* ── Columns ── */
[data-testid="stHorizontalBlock"] { align-items: flex-start !important; gap: 0 !important; }
[data-testid="column"] { padding: 0 5px !important; min-width: 0 !important; }
[data-testid="column"]:first-child { padding-left: 0 !important; }
[data-testid="column"]:last-child  { padding-right: 0 !important; }

/* ── Code ── */
.stCodeBlock, pre { border-radius: 7px !important; }

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] { border-bottom: 1px solid #e2e8f0 !important; }
.stTabs [aria-selected="true"] { border-bottom: 2px solid #2563eb !important; color: #2563eb !important; }
</style>""", unsafe_allow_html=True)


def _gap(px=16):
    st.markdown(f"<div style='height:{px}px'></div>", unsafe_allow_html=True)

def _label(text, mb=8):
    st.markdown(
        f"<p style='font-size:10px;font-weight:700;color:{C['sub']};letter-spacing:.09em;"
        f"text-transform:uppercase;margin:0 0 {mb}px;'>{text}</p>",
        unsafe_allow_html=True)

def _page_header(icon, title, subtitle=""):
    st.markdown(
        f"<div style='background:{C['s1']};border-bottom:1px solid {C['bd']};"
        f"box-shadow:0 1px 4px rgba(0,0,0,.06);position:sticky;top:0;z-index:100;'>"
        f"<div style='max-width:900px;margin:0 auto;padding:15px 24px 13px;"
        f"display:flex;align-items:center;gap:10px;'>"
        f"<span style='font-size:18px'>{icon}</span>"
        f"<div>"
        f"<div style='font-size:16px;font-weight:700;color:{C['tx']};'>{title}</div>"
        + (f"<div style='font-size:11px;color:{C['sub']};margin-top:1px;'>{subtitle}</div>" if subtitle else "")
        + "</div></div></div><div style='height:18px'></div>",
        unsafe_allow_html=True)

def _W():
    st.markdown(f"<div style='max-width:900px;margin:0 auto;padding:0 24px 40px;'>",
                unsafe_allow_html=True)

def _Wend():
    st.markdown("</div>", unsafe_allow_html=True)

def _info_box(text, bg=None, bdr=None, tc=None):
    bg  = bg  or C["bluelt"]
    bdr = bdr or C["bluebdr"]
    tc  = tc  or C["blue"]
    st.markdown(
        f"<div style='background:{bg};border:1px solid {bdr};border-left:4px solid {tc};"
        f"border-radius:8px;padding:12px 16px;margin-bottom:14px;'>"
        f"<p style='font-size:13px;color:{C['tx2']};margin:0;'>{text}</p></div>",
        unsafe_allow_html=True)

def _progress(val, color=None, label="", height=7):
    color = color or C["blue"]
    pct   = max(0, min(100, int(val)))
    bar_c = C["green"] if pct >= 80 else C["amber"] if pct >= 55 else C["red"]
    lbl   = (
        f"<div style='display:flex;justify-content:space-between;margin-bottom:5px;'>"
        f"<span style='font-size:11px;color:{C['sub']};font-weight:500;'>{label}</span>"
        f"<span style='font-size:11px;font-weight:700;color:{C['tx']};font-family:\"JetBrains Mono\",monospace;'>{pct}</span>"
        f"</div>"
    ) if label else ""
    st.markdown(
        f"{lbl}<div style='background:{C['s3']};border-radius:{height}px;height:{height}px;overflow:hidden;'>"
        f"<div style='width:{pct}%;height:100%;background:{bar_c};border-radius:{height}px;'></div></div>",
        unsafe_allow_html=True)

def _grade_col(g):
    return {"A":C["green"],"B":C["blue"],"C":C["amber"],"D":"#ea580c","F":C["red"]}.get(g,C["sub"])

def _rec_col(r):
    return {"Strong Hire":C["green"],"Hire":C["blue"],
            "Interview Further":C["amber"],"Decline":C["red"]}.get(r,C["sub"])

def _badge(text, color=None):
    c = color or C["blue"]
    return (f"<span style='background:{c}18;color:{c};border:1px solid {c}35;"
            f"padding:2px 9px;border-radius:5px;font-size:11px;font-weight:500;"
            f"display:inline-block;margin:2px 3px 2px 0;white-space:nowrap;'>{text}</span>")

def _rec_badge(r):
    c = _rec_col(r)
    return (f"<span style='background:{c}15;color:{c};border:1px solid {c}30;"
            f"padding:4px 12px;border-radius:20px;font-size:11px;font-weight:600;"
            f"white-space:nowrap;'>{r}</span>")

def _gmail_button(label, to, from_addr, subject, body, full_width=False):
    import urllib.parse
    url = (f"https://mail.google.com/mail/?view=cm&fs=1"
           f"&to={urllib.parse.quote(to)}"
           f"&from={urllib.parse.quote(from_addr)}"
           f"&su={urllib.parse.quote(subject)}"
           f"&body={urllib.parse.quote(body)}")
    w = "100%" if full_width else "auto"
    st.markdown(
        f"<a href='{url}' target='_blank' style='text-decoration:none;display:inline-block;width:{w};'>"
        f"<div style='background:{C['green']};color:#fff;border-radius:8px;"
        f"padding:11px 20px;font-size:13px;font-weight:600;text-align:center;"
        f"cursor:pointer;box-shadow:0 2px 6px rgba(22,163,74,.28);display:inline-block;width:{w};'>"
        f"✉️ {label}</div></a>"
        f"<p style='font-size:11px;color:{C['sub']};margin-top:6px;'>"
        f"Opens Gmail with a pre-filled draft — you can edit before sending.</p>",
        unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════
# PERSISTENT SIDEBAR  ─ renders on EVERY page (logged in or not)
# ══════════════════════════════════════════════════════════════════
def _sidebar():
    """
    Sidebar: expanded (260px) or collapsed (60px icon strip).
    Toggle stored in st.session_state.sidebar_open.
    NO query params — uses st.button for reliable reruns.
    NO help= on any button — prevents tooltip text popup.
    """
    if "sidebar_open" not in st.session_state:
        st.session_state.sidebar_open = True

    open_     = st.session_state.sidebar_open
    logged_in = st.session_state.get("logged_in", False)
    role      = st.session_state.get("role", "")
    username  = st.session_state.get("username", "")
    is_user   = (role == "User")
    nav       = st.session_state.get("nav", "Dashboard")

    # Width changes with state
    sb_w = "260px" if open_ else "62px"

    # Inject CSS — sidebar width + hide Streamlit's own collapse button
    st.markdown(f"""<style>
section[data-testid="stSidebar"] {{
    width: {sb_w} !important;
    min-width: {sb_w} !important;
    max-width: {sb_w} !important;
    transform: translateX(0px) !important;
    background-color: #ffffff !important;
    border-right: 2px solid #e2e8f0 !important;
    overflow: hidden !important;
    box-shadow: 2px 0 8px rgba(0,0,0,0.06) !important;
}}
section[data-testid="stSidebar"] > div:first-child {{
    width: {sb_w} !important;
    background-color: #ffffff !important;
    overflow-x: hidden !important;
    padding: 0 !important;
}}
section[data-testid="stSidebar"] [data-testid="stVerticalBlock"] {{
    background-color: #ffffff !important;
    padding: {"8px 10px" if open_ else "6px 4px"} !important;
    gap: {"2px" if not open_ else "0px"} !important;
}}
/* Hide Streamlit's own collapse arrow */
button[data-testid="baseButton-headerNoPadding"],
[data-testid="collapsedControl"],
[data-testid="stSidebarCollapseButton"] {{
    display: none !important;
}}
{"" if open_ else """
/* Collapsed: icon buttons square and centred */
section[data-testid="stSidebar"] .stButton > button {
    width: 46px !important;
    height: 44px !important;
    padding: 0 !important;
    margin: 1px auto !important;
    font-size: 20px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    background: transparent !important;
    border: none !important;
    border-radius: 8px !important;
    box-shadow: none !important;
    line-height: 1 !important;
}
section[data-testid="stSidebar"] .stButton > button:hover {
    background: #eff6ff !important;
    border: none !important;
    box-shadow: none !important;
}
"""}
</style>""", unsafe_allow_html=True)

    with st.sidebar:

        # ════════════════════════════════════════
        # EXPANDED
        # ════════════════════════════════════════
        if open_:
            # Brand row + collapse button
            col_brand, col_toggle = st.columns([5, 1])
            with col_brand:
                st.markdown(
                    f"<div style='padding:14px 0 6px;'>"
                    f"<div style='font-size:17px;font-weight:800;color:#2563eb;'>🛡️ DevTrust</div>"
                    f"<div style='font-size:10px;color:#6b7280;margin-top:1px;'>AI Developer Screening</div>"
                    f"</div>", unsafe_allow_html=True)
            with col_toggle:
                st.markdown("<div style='padding-top:14px;'>", unsafe_allow_html=True)
                if st.button("☰", key="sb_collapse"):
                    st.session_state.sidebar_open = False; st.rerun()
                st.markdown("</div>", unsafe_allow_html=True)

            st.divider()

            # ── Guest ──────────────────────────────────────
            if not logged_in:
                st.write("**Welcome!**")
                st.caption("Sign in or create an account.")
                if st.button("🔐  Sign In",       key="sb_signin",        use_container_width=True):
                    st.session_state.page = "login";    st.rerun()
                if st.button("✨  Create Account", key="sb_register",      use_container_width=True):
                    st.session_state.page = "register"; st.rerun()
                st.divider()
                if st.button("🏠  Home",           key="sb_home_guest",    use_container_width=True):
                    st.session_state.page = "home";     st.rerun()
                if st.button("🔒  Privacy",        key="sb_privacy_guest", use_container_width=True):
                    st.session_state.page = "privacy";  st.rerun()
                return None, None, None, None

            # ── Logged in ──────────────────────────────────
            all_creds = _all_credentials()
            display   = all_creds.get(username, {}).get("display", username)
            role_lbl  = "Developer" if is_user else "Organisation"
            role_clr  = "#16a34a" if is_user else "#2563eb"

            st.markdown(
                f"<div style='background:#f8fafc;border:1px solid #e2e8f0;border-radius:10px;"
                f"padding:10px 12px;margin-bottom:8px;display:flex;align-items:center;gap:10px;'>"
                f"<div style='width:30px;height:30px;border-radius:50%;background:{role_clr}20;"
                f"border:2px solid {role_clr}60;display:flex;align-items:center;justify-content:center;"
                f"font-size:14px;flex-shrink:0;'>{'👤' if is_user else '🏢'}</div>"
                f"<div><div style='font-size:13px;font-weight:700;color:#111827;'>{display}</div>"
                f"<div style='font-size:11px;color:{role_clr};font-weight:600;'>{role_lbl}</div>"
                f"</div></div>", unsafe_allow_html=True)

            # Notifications
            n_count = unread_count(username, role)
            with st.expander(f"🔔 Notifications{f' ({n_count})' if n_count else ''}", expanded=False):
                notifs = get_notifications(username, role)
                if not notifs:
                    st.caption("No notifications yet.")
                else:
                    for n in notifs[:6]:
                        st.caption(f"{'🔵 ' if not n.get('read') else ''}{n['msg']}")
                if n_count and st.button("✓ Mark all read", key="sb_mark_read", use_container_width=True):
                    mark_all_read(username, role); st.rerun()

            st.divider()
            st.markdown("<p style='font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.08em;text-transform:uppercase;margin:0 0 6px;'>Navigation</p>", unsafe_allow_html=True)

            nav_items = (
                [("🏠","Home","Home"),("📊","Dashboard","Dashboard"),
                 ("📤","Upload","Upload & Score"),("📈","Results","Trust Report"),
                 ("🗂️","Files","File Manager")]
                if is_user else
                [("🏠","Home","Home"),("📊","Dashboard","Dashboard"),
                 ("🏆","Portfolio","Portfolio"),("📋","Detail","Project Detail")]
            )
            for icon, key, label in nav_items:
                if key != "Home" and nav == key:
                    st.markdown(
                        f"<div style='background:#eff6ff;border:1px solid #bfdbfe;border-radius:8px;"
                        f"padding:9px 12px;margin:2px 0;display:flex;align-items:center;gap:9px;'>"
                        f"<span style='font-size:16px;'>{icon}</span>"
                        f"<span style='font-size:13px;font-weight:700;color:#2563eb;'>{label}</span>"
                        f"</div>", unsafe_allow_html=True)
                else:
                    if st.button(f"{icon}  {label}", key=f"sb_nav_{key}", use_container_width=True):
                        if key == "Home":
                            for k in list(st.session_state.keys()): del st.session_state[k]
                            st.rerun()
                        else:
                            st.session_state.nav = key; st.rerun()

            st.divider()
            st.markdown("<p style='font-size:10px;font-weight:700;color:#6b7280;letter-spacing:.08em;text-transform:uppercase;margin:0 0 6px;'>LLM Engine</p>", unsafe_allow_html=True)
            q_ok = bool(GROQ_API_KEY); g_ok = bool(GOOGLE_API_KEY)
            avail = (["Groq"] if q_ok else []) + (["Google Gemini"] if g_ok else []) or ["Groq","Google Gemini"]
            st.caption(("🟢 Groq  " if q_ok else "🔴 Groq  ") + ("🟢 Gemini" if g_ok else "⚫ Gemini"))
            provider = st.selectbox("Provider", avail, label_visibility="collapsed", key="sb_provider")
            models   = (["llama-3.3-70b-versatile","llama-3.1-8b-instant","gemma2-9b-it"]
                        if provider=="Groq" else
                        ["gemini-2.0-flash","gemini-2.0-flash-lite","gemini-1.5-flash-latest"])
            api_key  = GROQ_API_KEY if provider=="Groq" else GOOGLE_API_KEY
            model    = st.selectbox("Model", models, label_visibility="collapsed", key="sb_model")
            st.divider()
            if st.button("🚪  Sign Out", key="sb_signout", use_container_width=True):
                for k in list(st.session_state.keys()): del st.session_state[k]
                st.rerun()

        # ════════════════════════════════════════
        # COLLAPSED — emoji icon strip
        # ════════════════════════════════════════
        else:
            # Expand button — NO help= param (prevents tooltip text)
            if st.button("☰", key="sb_expand"):
                st.session_state.sidebar_open = True; st.rerun()

            st.divider()

            if not logged_in:
                for ico, page in [("🔐","login"),("✨","register"),("🏠","home"),("🔒","privacy")]:
                    if st.button(ico, key=f"sb_c_{page}"):
                        st.session_state.page = page; st.rerun()
                return None, None, None, None

            # Role chip
            role_clr = "#16a34a" if is_user else "#2563eb"
            st.markdown(
                f"<div style='width:36px;height:36px;border-radius:50%;"
                f"background:{role_clr}20;border:2px solid {role_clr}60;"
                f"display:flex;align-items:center;justify-content:center;"
                f"font-size:16px;margin:4px auto;'>{'👤' if is_user else '🏢'}</div>",
                unsafe_allow_html=True)

            st.divider()

            # Nav icons — NO help= so no tooltip text appears
            nav_items_c = (
                [("🏠","Home"),("📊","Dashboard"),("📤","Upload"),("📈","Results"),("🗂️","Files")]
                if is_user else
                [("🏠","Home"),("📊","Dashboard"),("🏆","Portfolio"),("📋","Detail")]
            )
            for ico, key in nav_items_c:
                is_active = (key != "Home" and nav == key)
                if is_active:
                    st.markdown(
                        f"<div style='width:44px;height:44px;margin:2px auto;"
                        f"background:#eff6ff;border:2px solid #bfdbfe;border-radius:10px;"
                        f"display:flex;align-items:center;justify-content:center;"
                        f"font-size:20px;'>{ico}</div>",
                        unsafe_allow_html=True)
                else:
                    if st.button(ico, key=f"sb_c_nav_{key}"):
                        if key == "Home":
                            for k in list(st.session_state.keys()): del st.session_state[k]
                            st.rerun()
                        else:
                            st.session_state.nav = key; st.rerun()

            st.divider()

            # Notifications icon
            n_count = unread_count(username, role)
            if st.button("🔴" if n_count else "🔔", key="sb_c_notif"):
                st.session_state.sidebar_open = True; st.rerun()

            # Sign out icon
            if st.button("🚪", key="sb_c_out"):
                for k in list(st.session_state.keys()): del st.session_state[k]
                st.rerun()

            # LLM defaults
            q_ok = bool(GROQ_API_KEY); g_ok = bool(GOOGLE_API_KEY)
            avail    = (["Groq"] if q_ok else []) + (["Google Gemini"] if g_ok else []) or ["Groq","Google Gemini"]
            provider = avail[0]
            models   = ["llama-3.3-70b-versatile"] if provider=="Groq" else ["gemini-2.0-flash"]
            api_key  = GROQ_API_KEY if provider=="Groq" else GOOGLE_API_KEY
            model    = models[0]

    return (st.session_state.get("sb_provider", provider),
            api_key,
            st.session_state.get("sb_model", model),
            nav)


# ══════════════════════════════════════════════════════════
# PUBLIC — HOME
# ══════════════════════════════════════════════════════════
def render_home():
    st.markdown(
        f"<div style='max-width:900px;margin:0 auto;padding:0 24px;'>",
        unsafe_allow_html=True)
    _gap(40)
    st.markdown(
        f"<div style='display:inline-flex;align-items:center;gap:6px;"
        f"background:{C['bluelt']};border:1px solid {C['bluebdr']};border-radius:20px;"
        f"padding:4px 14px;margin-bottom:20px;'>"
        f"<span style='width:6px;height:6px;border-radius:50%;background:{C['blue']};display:inline-block;'></span>"
        f"<span style='font-size:11px;color:{C['blue']};font-weight:600;letter-spacing:.07em;text-transform:uppercase;'>"
        f"AI-Powered Developer Due Diligence</span></div>",
        unsafe_allow_html=True)
    st.markdown(
        f"<h1 style='font-size:44px;font-weight:700;color:{C['tx']};line-height:1.1;"
        f"letter-spacing:-.03em;margin:0 0 16px;'>"
        f"The Trust Layer<br><span style='color:{C['blue']};'>for Developer Portfolios</span></h1>"
        f"<p style='font-size:15px;color:{C['sub']};max-width:500px;line-height:1.7;margin:0 0 28px;'>"
        f"Upload a project, pick its type, get an AI trust score in seconds. "
        f"Share a <strong style='color:{C['tx']};'>sanitized recruiter portal</strong> — "
        f"your IP never leaves your control.</p>",
        unsafe_allow_html=True)

    c1, c2, _ = st.columns([1, 1, 4])
    with c1:
        if st.button("🔐  Sign In", key="home_signin", use_container_width=True, type="primary"):
            st.session_state.page = "login"; st.rerun()
    with c2:
        if st.button("✨  Create Account", key="home_register", use_container_width=True):
            st.session_state.page = "register"; st.rerun()

    _gap(40)
    cols = st.columns(5)
    for col, (val, lbl, clr) in zip(cols, [
        ("100%","IP Protected",C["blue"]),("15+","File Formats",C["tx"]),
        ("20","Sample Projects",C["green"]),("6","Project Types",C["violet"]),
        ("<30s","Scoring Time",C["amber"]),
    ]):
        with col:
            st.markdown(
                f"<div style='background:{C['s1']};border:1px solid {C['bd']};border-radius:10px;"
                f"padding:16px;text-align:center;box-shadow:0 1px 3px rgba(0,0,0,.06);'>"
                f"<div style='font-size:24px;font-weight:700;color:{clr};'>{val}</div>"
                f"<div style='font-size:11px;color:{C['sub']};margin-top:4px;font-weight:500;'>{lbl}</div>"
                f"</div>", unsafe_allow_html=True)
    _gap(36)

    feats = [
        ("🔒","IP Sanitization",C["blue"],"Secrets and API keys auto-redacted before AI analysis."),
        ("🧠","AI Trust Score",C["violet"],"9-dimension weighted scoring adapted to your project type."),
        ("🏆","Ranked Portfolio",C["green"],"Organisations see projects ranked by score. No raw code exposed."),
        ("✉️","One-Click Contact",C["amber"],"Gmail pre-filled with candidate address and personalised message."),
        ("📁","Any Format",C["blue"],"ZIP, code, notebooks, reports, PDFs — upload as-is."),
        ("🔔","Notifications",C["red"],"Orgs notified on new submissions. Developers get confirmation."),
    ]
    r1 = st.columns(3, gap="medium")
    r2 = st.columns(3, gap="medium")
    for i, (ico, ttl, clr, desc) in enumerate(feats):
        with (r1 if i < 3 else r2)[i % 3]:
            st.markdown(
                f"<div style='background:{C['s1']};border:1px solid {C['bd']};border-radius:10px;"
                f"padding:18px;box-shadow:0 1px 3px rgba(0,0,0,.06);margin-bottom:10px;'>"
                f"<div style='width:34px;height:34px;border-radius:8px;background:{clr}15;"
                f"display:flex;align-items:center;justify-content:center;font-size:16px;margin-bottom:10px;'>{ico}</div>"
                f"<div style='font-weight:600;font-size:13px;color:{C['tx']};margin-bottom:5px;'>{ttl}</div>"
                f"<div style='font-size:12px;color:{C['sub']};line-height:1.6;'>{desc}</div></div>",
                unsafe_allow_html=True)
    st.markdown("</div>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# PUBLIC — LOGIN
# ══════════════════════════════════════════════════════════
def render_login():
    """Full-width login page — sidebar is hidden when this renders."""
    # Page background + centred card
    st.markdown(f"""<style>
.stApp, [data-testid="stAppViewContainer"],
[data-testid="stMain"], .main {{
    background: {C['bg']} !important;
}}
</style>""", unsafe_allow_html=True)

    _gap(40)

    # Back to home
    _, top_mid, _ = st.columns([1, 1.4, 1])
    with top_mid:
        if st.button("← Back to Home", key="li_back_home", use_container_width=False):
            st.session_state.page = "home"; st.rerun()

    _gap(16)

    _, mid, _ = st.columns([1, 1.4, 1])
    with mid:
        st.markdown(
            f"<div style='text-align:center;margin-bottom:24px;'>"
            f"<div style='font-size:32px;font-weight:800;color:{C['blue']};margin-bottom:6px;'>"
            f"🛡️ DevTrustProtocol</div>"
            f"<div style='font-size:14px;color:{C['sub']};'>Sign in to your account</div></div>",
            unsafe_allow_html=True)

        st.markdown(
            f"<div style='background:{C['s1']};border:1px solid {C['bd']};"
            f"border-radius:14px;padding:32px;box-shadow:0 4px 24px rgba(0,0,0,.10);'>",
            unsafe_allow_html=True)

        username = st.text_input("Username", placeholder="Enter your username", key="li_user")
        _gap(4)
        password = st.text_input("Password", type="password", placeholder="Enter your password", key="li_pass")
        _gap(14)

        if st.button("🔐  Sign In", key="li_submit", use_container_width=True, type="primary"):
            u = (username or "").strip()
            p = (password or "").strip()
            all_creds = _all_credentials()
            if u in all_creds and all_creds[u]["password"] == p:
                st.session_state.update({
                    "logged_in": True, "username": u,
                    "role":     all_creds[u]["role"],
                    "display_name": all_creds[u].get("display", u),
                    "page": "app", "dark_mode": False,
                    "nav": "Dashboard",
                })
                st.rerun()
            else:
                st.error("❌ Invalid username or password.")

        _gap(10)
        st.markdown(
            f"<div style='text-align:center;font-size:13px;color:{C['sub']};margin-bottom:8px;'>"
            f"Don't have an account?</div>",
            unsafe_allow_html=True)
        if st.button("✨  Create Account", key="li_to_register", use_container_width=True):
            st.session_state.page = "register"; st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)
        _gap(12)

        # Demo credentials
        st.markdown(
            f"<div style='background:{C['s2']};border:1px solid {C['bd']};"
            f"border-radius:10px;padding:14px 16px;'>"
            f"<p style='font-size:10px;font-weight:700;color:{C['sub']};letter-spacing:.08em;"
            f"text-transform:uppercase;margin:0 0 10px;'>Demo Accounts</p>"
            f"<div style='font-size:13px;color:{C['sub']};line-height:2;'>"
            f"👤 <strong style='color:{C['tx']};'>alex_dev</strong>"
            f"<span style='font-family:monospace;background:{C['s3']};padding:2px 8px;border-radius:4px;margin-left:8px;font-size:12px;'>dev@2025!</span>"
            f"<span style='color:{C['green']};font-size:11px;margin-left:8px;'>Developer</span><br>"
            f"🏢 <strong style='color:{C['tx']};'>techcorp_hr</strong>"
            f"<span style='font-family:monospace;background:{C['s3']};padding:2px 8px;border-radius:4px;margin-left:8px;font-size:12px;'>tc#hire25</span>"
            f"<span style='color:{C['blue']};font-size:11px;margin-left:8px;'>Organisation</span>"
            f"</div></div>",
            unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# PUBLIC — REGISTER
# ══════════════════════════════════════════════════════════
def render_register():
    _gap(48)
    _, mid, _ = st.columns([1, 1.2, 1])
    with mid:
        st.markdown(
            f"<div style='text-align:center;margin-bottom:22px;'>"
            f"<div style='font-size:28px;font-weight:800;color:{C['blue']};margin-bottom:4px;'>🛡️ DevTrustProtocol</div>"
            f"<div style='font-size:13px;color:{C['sub']};'>Create your account</div></div>",
            unsafe_allow_html=True)
        st.markdown(
            f"<div style='background:{C['s1']};border:1px solid {C['bd']};"
            f"border-radius:12px;padding:28px;box-shadow:0 4px 20px rgba(0,0,0,.09);'>",
            unsafe_allow_html=True)

        display  = st.text_input("Full Name / Display Name", placeholder="e.g. Alex Chen", key="reg_display")
        username = st.text_input("Username", placeholder="e.g. alex_dev  (no spaces)", key="reg_user")
        email    = st.text_input("Email (optional)", placeholder="your@email.com", key="reg_email")
        role_sel = st.radio("Account Type", ["Developer", "Organisation"],
                            horizontal=True, key="reg_role")
        _gap(4)
        password  = st.text_input("Password", type="password", placeholder="Min 6 characters", key="reg_pass")
        password2 = st.text_input("Confirm Password", type="password", placeholder="Repeat password", key="reg_pass2")
        _gap(10)

        if st.button("✨ Create Account", key="reg_submit", use_container_width=True, type="primary"):
            if password != password2:
                st.error("❌ Passwords do not match.")
            else:
                role = "User" if role_sel == "Developer" else "Recruiter"
                ok, msg = register_user(
                    username=(username or "").strip(),
                    password=password,
                    display=(display or "").strip(),
                    role=role,
                    email=(email or "").strip(),
                )
                if ok:
                    st.success(f"✅ {msg} You can now sign in.")
                    _gap(6)
                    if st.button("→ Go to Sign In", key="reg_goto_login", use_container_width=True, type="primary"):
                        st.session_state.page = "login"; st.rerun()
                else:
                    st.error(f"❌ {msg}")

        _gap(6)
        st.markdown(
            f"<div style='text-align:center;font-size:12px;color:{C['sub']};'>"
            f"Already have an account?</div>",
            unsafe_allow_html=True)
        if st.button("Sign In →", key="reg_to_login", use_container_width=True):
            st.session_state.page = "login"; st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)
        _gap(10)
        _info_box(
            "🔒 <strong>Developer</strong> accounts can upload projects, score them, and manage submissions. "
            "<strong>Organisation</strong> accounts can view the ranked portfolio and contact candidates.",
            C["bluelt"], C["bluebdr"], C["blue"])


# ══════════════════════════════════════════════════════════
# PUBLIC — PRIVACY
# ══════════════════════════════════════════════════════════
def render_privacy():
    _gap(32)
    _, mc, _ = st.columns([1, 3, 1])
    with mc:
        if st.button("← Back", key="priv_back", use_container_width=True):
            st.session_state.page = "home"; st.rerun()
        _gap(12)
        st.markdown(
            f"<div style='background:{C['s1']};border:1px solid {C['bd']};"
            f"border-radius:12px;padding:30px;box-shadow:0 1px 4px rgba(0,0,0,.07);'>",
            unsafe_allow_html=True)
        st.markdown(
            f"<h1 style='font-size:20px;font-weight:700;color:{C['tx']};margin-bottom:4px;'>🔒 Privacy Policy</h1>"
            f"<p style='font-size:12px;color:{C['sub']};margin-bottom:22px;'>DevTrustProtocol v3.0 · 2025</p>",
            unsafe_allow_html=True)
        for title, body in [
            ("1. What We Collect","Only what you explicitly upload: project files, resume documents, and profile metadata."),
            ("2. IP Protection","All source code sanitized before any AI call. Recruiters cannot view raw files, credentials, or source code under any circumstances."),
            ("3. AI Processing","Sanitized project data sent to your configured AI provider using your own API keys only."),
            ("4. Recruiter Access","Organisations see: project name, type, description, trust score, grade, working %, and recommendation — never source code."),
            ("5. Your Rights","You may delete any of your submissions at any time from your dashboard."),
        ]:
            st.markdown(
                f"<div style='margin-bottom:16px;padding-bottom:16px;border-bottom:1px solid {C['bd']};'>"
                f"<h3 style='font-size:13px;font-weight:700;color:{C['tx']};margin-bottom:5px;'>{title}</h3>"
                f"<div style='font-size:13px;color:{C['sub']};line-height:1.7;'>{body}</div></div>",
                unsafe_allow_html=True)
        _info_box("🛡️ You are in full control. Recruiters cannot see your raw code or credentials.")
        st.markdown("</div>", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════
# DEVELOPER DASHBOARD
# ══════════════════════════════════════════════════════════
def render_user_dashboard(username):
    _page_header("📊", f"Welcome, {_all_credentials().get(username,{}).get('display',username)}", "Your developer activity overview")
    _W()
    subs   = [s for s in get_all_submissions() if s.get("uploaded_by") == username]
    total  = len(subs)
    avg_sc = round(sum(s["result"].get("trust_score",0) for s in subs)/total) if total else 0
    best   = max((s["result"].get("trust_score",0) for s in subs), default=0)
    top_gr = subs[0]["result"].get("grade","—") if subs else "—"
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Projects Submitted", total)
    c2.metric("Average Score", f"{avg_sc}/100" if total else "—")
    c3.metric("Best Score", f"{best}/100" if total else "—")
    c4.metric("Top Grade", top_gr)
    _gap(22)

    result = st.session_state.get("result")
    if result:
        score = result.get("trust_score",0); grade = result.get("grade","?")
        rec   = result.get("recommendation",""); ptype = result.get("project_type","")
        gc    = _grade_col(grade); ptc = TYPE_COLOR.get(ptype, C["sub"])
        st.markdown(
            f"<div style='background:{C['s1']};border:1px solid {C['bd']};border-radius:10px;"
            f"padding:16px 18px;box-shadow:0 1px 3px rgba(0,0,0,.07);margin-bottom:14px;'>"
            f"<div style='display:flex;align-items:center;justify-content:space-between;flex-wrap:wrap;gap:8px;'>"
            f"<div>"
            f"<p style='font-size:10px;font-weight:700;color:{C['sub']};letter-spacing:.08em;text-transform:uppercase;margin:0 0 3px;'>Active Session</p>"
            f"<p style='font-size:14px;font-weight:600;color:{C['tx']};margin:0 0 4px;'>{result.get('source','Project')}</p>"
            + (f"<span style='background:{ptc}15;color:{ptc};border:1px solid {ptc}30;border-radius:10px;padding:1px 8px;font-size:10px;font-weight:600;'>{ptype}</span>" if ptype else "")
            + f"</div>"
            f"<div style='display:flex;gap:14px;align-items:center;'>"
            f"<div style='text-align:center;'><div style='font-size:24px;font-weight:700;color:{gc};'>{score}</div><div style='font-size:10px;color:{C['sub']};'>Score</div></div>"
            f"<div style='text-align:center;'><div style='font-size:24px;font-weight:700;color:{gc};'>{grade}</div><div style='font-size:10px;color:{C['sub']};'>Grade</div></div>"
            f"<div>{_rec_badge(rec)}</div></div></div></div>",
            unsafe_allow_html=True)
        b1,b2,_ = st.columns([1,1,4])
        with b1:
            if st.button("📈 View Report", key="dash_view_report", use_container_width=True, type="primary"):
                st.session_state.nav = "Results"; st.rerun()
        with b2:
            if st.button("📤 Score New", key="dash_score_new", use_container_width=True):
                st.session_state.nav = "Upload"; st.rerun()
    else:
        _info_box("No project scored yet. Head to <strong>Upload & Score</strong> to get started.")
        if st.button("📤 Upload & Score →", key="dash_goto_upload", type="primary"):
            st.session_state.nav = "Upload"; st.rerun()

    _gap(20)
    if subs:
        _label("Your Submission History")
        _gap(6)
        for sub in subs[:10]:
            r    = sub.get("result",{}); p = sub.get("profile",{})
            sc   = r.get("trust_score",0); gr = r.get("grade","?")
            rec  = r.get("recommendation",""); gc = _grade_col(gr)
            proj = r.get("source",sub.get("_key","Project"))
            ptype= sub.get("project_type",""); pdesc = sub.get("description","").strip()
            saved= sub.get("saved_at",""); key = sub.get("_key","")
            ptc  = TYPE_COLOR.get(ptype, C["sub"])
            st.markdown(
                f"<div style='background:{C['s1']};border:1px solid {C['bd']};border-radius:9px;"
                f"padding:13px 15px;box-shadow:0 1px 3px rgba(0,0,0,.05);margin-bottom:7px;'>"
                f"<div style='display:flex;align-items:flex-start;justify-content:space-between;flex-wrap:wrap;gap:8px;'>"
                f"<div style='flex:1;min-width:160px;'>"
                f"<div style='display:flex;align-items:center;gap:6px;margin-bottom:2px;'>"
                f"<span style='font-size:13px;font-weight:600;color:{C['tx']};'>📁 {proj}</span>"
                + (f"<span style='background:{ptc}15;color:{ptc};border:1px solid {ptc}30;border-radius:10px;padding:1px 7px;font-size:10px;font-weight:600;'>{ptype}</span>" if ptype else "")
                + f"</div>"
                f"<div style='font-size:11px;color:{C['sub']};'>{saved}</div>"
                + (f"<div style='font-size:11px;color:{C['sub']};margin-top:3px;'>{pdesc[:80]}{'…' if len(pdesc)>80 else ''}</div>" if pdesc else "")
                + f"</div>"
                f"<div style='display:flex;align-items:center;gap:10px;'>"
                f"<span style='font-size:17px;font-weight:700;color:{gc};'>{sc}/100 {gr}</span>"
                f"{_rec_badge(rec)}</div></div>"
                f"<div style='display:flex;flex-wrap:wrap;gap:3px;margin-top:7px;'>"
                + "".join(_badge(t) for t in r.get("tech_stack_summary",[])[:5])
                + "</div></div>",
                unsafe_allow_html=True)
            d1,d2,_ = st.columns([1,1,5])
            with d1:
                if st.button("📈 View", key=f"dh_view_{key}", use_container_width=True):
                    st.session_state.result         = r
                    st.session_state.resume_profile = p
                    st.session_state.resume_text    = sub.get("resume_text","")
                    st.session_state.nav = "Results"; st.rerun()
            with d2:
                if st.button("🗑️ Delete", key=f"dh_del_{key}", use_container_width=True):
                    st.session_state[f"cdel_{key}"] = True; st.rerun()
            if st.session_state.get(f"cdel_{key}"):
                st.warning(f"Delete **{proj}**? Cannot be undone.")
                y,n,_ = st.columns([1,1,5])
                with y:
                    if st.button("✅ Confirm", key=f"dh_yes_{key}", use_container_width=True):
                        delete_submission(key); st.session_state.pop(f"cdel_{key}",None)
                        st.success("Deleted."); st.rerun()
                with n:
                    if st.button("❌ Cancel", key=f"dh_no_{key}", use_container_width=True):
                        st.session_state.pop(f"cdel_{key}",None); st.rerun()
    else:
        st.markdown(
            f"<div style='text-align:center;padding:36px;background:{C['s1']};"
            f"border:1px solid {C['bd']};border-radius:10px;'>"
            f"<div style='font-size:28px;margin-bottom:8px;'>📭</div>"
            f"<div style='font-size:13px;font-weight:600;color:{C['tx']};margin-bottom:5px;'>No submissions yet</div>"
            f"<div style='font-size:12px;color:{C['sub']};'>Score your first project to see it here.</div></div>",
            unsafe_allow_html=True)
    _Wend()


# ══════════════════════════════════════════════════════════
# ORG DASHBOARD
# ══════════════════════════════════════════════════════════
def render_org_dashboard(username):
    org_email = _all_credentials().get(username,{}).get("email","")
    display   = _all_credentials().get(username,{}).get("display",username)
    _page_header("📊", f"Welcome, {display}", "Organisation recruiter overview")
    _W()
    real_subs = get_all_submissions()
    all_p = []
    for sub in real_subs:
        r = sub.get("result",{})
        all_p.append({"score":r.get("trust_score",0),"grade":r.get("grade","?"),
                      "rec":r.get("recommendation",""),"_sub":sub})
    for gh in GITHUB_PROJECTS:
        all_p.append({"score":gh["score"],"grade":gh["grade"],"rec":gh["recommendation"],"_sub":None})
    all_p.sort(key=lambda x: x["score"], reverse=True)
    scores = [p["score"] for p in all_p]
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Total Projects",   len(all_p))
    c2.metric("Real Submissions", len(real_subs))
    c3.metric("Avg Score",        f"{round(sum(scores)/len(scores))}/100" if scores else "—")
    c4.metric("Strong Hire",      sum(1 for p in all_p if p["rec"]=="Strong Hire"))
    _gap(20)

    grades_cnt = {}
    for p in all_p: grades_cnt[p["grade"]] = grades_cnt.get(p["grade"],0)+1
    _label("Grade Distribution")
    _gap(6)
    gcols = st.columns(5)
    for col,g in zip(gcols,["A","B","C","D","F"]):
        cnt = grades_cnt.get(g,0); gc = _grade_col(g)
        with col:
            st.markdown(
                f"<div style='background:{C['s1']};border:1px solid {C['bd']};border-radius:9px;"
                f"padding:12px;text-align:center;box-shadow:0 1px 3px rgba(0,0,0,.05);'>"
                f"<div style='font-size:20px;font-weight:700;color:{gc};'>{g}</div>"
                f"<div style='font-size:16px;font-weight:700;color:{C['tx']};margin-top:3px;'>{cnt}</div>"
                f"<div style='font-size:10px;color:{C['sub']};margin-top:1px;'>projects</div></div>",
                unsafe_allow_html=True)
    _gap(20)

    if real_subs:
        _label("Top Candidates")
        _gap(6)
        for sub in real_subs[:5]:
            r    = sub.get("result",{}); p = sub.get("profile",{})
            sc   = r.get("trust_score",0); gr = r.get("grade","?")
            rec  = r.get("recommendation",""); gc = _grade_col(gr)
            name = p.get("name","Unknown"); proj = r.get("source","Project")
            email= p.get("email",""); key = sub.get("_key","")
            ptype= sub.get("project_type",""); pdesc = sub.get("description","").strip()
            ptc  = TYPE_COLOR.get(ptype, C["sub"])
            st.markdown(
                f"<div style='background:{C['s1']};border:1px solid {C['bd']};border-radius:9px;"
                f"padding:13px 15px;box-shadow:0 1px 3px rgba(0,0,0,.05);margin-bottom:7px;'>"
                f"<div style='display:flex;align-items:flex-start;justify-content:space-between;flex-wrap:wrap;gap:8px;'>"
                f"<div style='flex:1;min-width:160px;'>"
                f"<div style='display:flex;align-items:center;gap:6px;margin-bottom:2px;'>"
                f"<span style='font-size:13px;font-weight:600;color:{C['tx']};'>👤 {name}</span>"
                + (f"<span style='background:{ptc}15;color:{ptc};border:1px solid {ptc}30;border-radius:10px;padding:1px 7px;font-size:10px;font-weight:600;'>{ptype}</span>" if ptype else "")
                + f"</div>"
                f"<div style='font-size:11px;color:{C['sub']};'>📁 {proj}</div>"
                + (f"<div style='font-size:11px;color:{C['sub']};margin-top:3px;'>{pdesc[:100]}{'…' if len(pdesc)>100 else ''}</div>" if pdesc else "")
                + f"</div>"
                f"<div style='display:flex;align-items:center;gap:10px;'>"
                f"<span style='font-size:18px;font-weight:700;color:{gc};'>{sc}/100 {gr}</span>"
                f"{_rec_badge(rec)}</div></div></div>",
                unsafe_allow_html=True)
            d1, d2, _ = st.columns([1, 1, 5])
            with d1:
                if st.button("📋 Detail", key=f"od_detail_{key}", use_container_width=True):
                    st.session_state.detail_real = sub
                    st.session_state.detail_gh   = None
                    st.session_state.nav = "Detail"; st.rerun()
            with d2:
                if email:
                    import urllib.parse as _up
                    subject = "Re: Your Project Submission — DevTrustProtocol"
                    body    = (f"Hi {name},\n\nI reviewed your project '{proj}' on DevTrustProtocol.\n\n"
                               f"Trust Score: {sc}/100 | Grade: {gr} | {rec}\n\n"
                               f"I'd love to connect. Are you available for a call this week?\n\nBest regards,")
                    gmail_url = (f"https://mail.google.com/mail/?view=cm&fs=1"
                                 f"&to={_up.quote(email)}"
                                 f"&from={_up.quote(org_email)}"
                                 f"&su={_up.quote(subject)}"
                                 f"&body={_up.quote(body)}")
                    st.markdown(
                        f"<a href='{gmail_url}' target='_blank' style='text-decoration:none;'>"
                        f"<button style='"
                        f"width:100%;background:#16a34a;color:#fff;border:none;"
                        f"border-radius:7px;padding:8px 0;font-size:13px;font-weight:600;"
                        f"cursor:pointer;font-family:Inter,sans-serif;"
                        f"display:flex;align-items:center;justify-content:center;gap:6px;"
                        f"'>📧 Contact</button></a>",
                        unsafe_allow_html=True)
                else:
                    st.button("📧 Contact", key=f"od_no_email_{key}",
                              disabled=True, use_container_width=True)
    else:
        _info_box("No real submissions yet. Ask developers to score their projects.")
    _gap(12)
    if st.button("🏆 View Full Portfolio →", key="od_portfolio", type="primary"):
        st.session_state.nav = "Portfolio"; st.rerun()
    _Wend()


# ══════════════════════════════════════════════════════════
# USER — UPLOAD & SCORE
# ══════════════════════════════════════════════════════════
def render_upload(provider, api_key, model):
    _page_header("📤","Upload & Score","Describe your project, choose type, upload files, then score.")
    _W()
    b1,b2,b3,_ = st.columns([1,1,1,4])
    with b1:
        if st.button("🏠 Dashboard", key="up_dash", use_container_width=True):
            st.session_state.nav = "Dashboard"; st.rerun()
    with b2:
        if st.session_state.get("result"):
            if st.button("📈 Results", key="up_results", use_container_width=True):
                st.session_state.nav = "Results"; st.rerun()
    with b3:
        if st.session_state.get("project_data"):
            if st.button("🗂️ Files", key="up_files", use_container_width=True):
                st.session_state.nav = "Files"; st.rerun()
    _gap(14)

    shared = load_shared_state()
    if shared:
        c1,c2 = st.columns([5,1])
        with c1: st.success(f"📡 Portal live — last scored: `{shared.get('saved_at','')}`")
        with c2:
            if st.button("🗑️ Reset", key="up_reset", use_container_width=True):
                clear_shared_state()
                for k in ["result","resume_profile","resume_text","project_data","sanitized_data","resume_bytes","resume_ext","resume_filename"]:
                    st.session_state.pop(k,None)
                st.rerun()
    else:
        st.info("📡 Organisation portal empty — score a project to publish it.")
    _gap(16)

    # Step 1
    st.markdown(
        f"<div style='background:{C['bluelt']};border:1px solid {C['bluebdr']};"
        f"border-left:4px solid {C['blue']};border-radius:8px;padding:14px 16px;margin-bottom:14px;'>"
        f"<p style='font-size:11px;font-weight:700;color:{C['blue']};letter-spacing:.07em;text-transform:uppercase;margin:0 0 12px;'>Step 1 — Project Info</p>",
        unsafe_allow_html=True)
    m1,m2 = st.columns([1,1], gap="large")
    with m1:
        _label("Project Name")
        st.text_input("pn", placeholder="e.g. Customer Churn Predictor", label_visibility="collapsed", key="pu_name")
    with m2:
        _label("Project Type")
        st.selectbox("pt", TrustScorer.PROJECT_TYPES, label_visibility="collapsed", key="pu_type")
    _gap(8)
    _label("Project Description  (shown to recruiters)")
    st.text_area("pd", placeholder="Describe what this project does, the problem it solves, key methods and results...",
                 height=90, label_visibility="collapsed", key="pu_desc")
    st.markdown("</div>", unsafe_allow_html=True)

    type_hints = {
        "Software Development":("🛠️","Include architecture decisions, tech stack, and how to run the project."),
        "Research & Analytics":("🔬","Describe your research question, data sources, methodology, and key findings."),
        "Data Science / ML":   ("🤖","Mention the ML task, dataset, model(s) used, evaluation metrics, and results."),
        "Design & UX":         ("🎨","Describe the design challenge, process, user research done, and prototype details."),
        "Business / Strategy": ("📊","State the business problem, your analytical approach, and the recommendations."),
        "Other / General":     ("📁","Briefly describe the project goals, approach, and outcomes."),
    }
    ptype_cur = st.session_state.get("pu_type","Software Development")
    hint_ico, hint_txt = type_hints.get(ptype_cur,("📁","Describe your project goals and approach."))
    st.markdown(
        f"<div style='background:{C['s2']};border:1px solid {C['bd']};border-radius:7px;"
        f"padding:9px 13px;margin-bottom:14px;display:flex;gap:9px;align-items:flex-start;'>"
        f"<span style='font-size:15px;'>{hint_ico}</span>"
        f"<p style='font-size:12px;color:{C['sub']};margin:0;line-height:1.6;'>"
        f"<strong style='color:{C['tx']};'>Tip:</strong> {hint_txt}</p></div>",
        unsafe_allow_html=True)

    # Step 2
    st.markdown(
        f"<div style='background:{C['s1']};border:1px solid {C['bd']};"
        f"border-left:4px solid {C['green']};border-radius:8px;padding:14px 16px;margin-bottom:14px;'>"
        f"<p style='font-size:11px;font-weight:700;color:{C['green']};letter-spacing:.07em;text-transform:uppercase;margin:0 0 12px;'>Step 2 — Upload Files</p>",
        unsafe_allow_html=True)
    uc1,uc2 = st.columns(2, gap="large")
    with uc1:
        _label("Project Files")
        st.markdown(f"<p style='font-size:12px;color:{C['sub']};margin-bottom:7px;'>ZIP, code, notebooks, reports, PDFs</p>",unsafe_allow_html=True)
        project_files = st.file_uploader("pf", type=ACCEPTED_PROJECT_TYPES, accept_multiple_files=True, label_visibility="collapsed", key="pu_p")
    with uc2:
        _label("Resume")
        st.markdown(f"<p style='font-size:12px;color:{C['sub']};margin-bottom:7px;'>PDF, DOCX, or JSON</p>",unsafe_allow_html=True)
        resume_file = st.file_uploader("rf", type=["pdf","docx","json"], accept_multiple_files=False, label_visibility="collapsed", key="pu_r")
    st.markdown("</div>", unsafe_allow_html=True)

    for k,v in [("project_data",None),("sanitized_data",None),("resume_text",""),
                ("resume_profile",{}),("result",None),("prev_names",[]),
                ("san_report",None),("resume_bytes",None),("resume_ext",".txt"),("resume_filename","resume")]:
        if k not in st.session_state: st.session_state[k] = v

    if project_files:
        names = sorted(f.name for f in project_files)
        if names != st.session_state.prev_names:
            for f in project_files: f.seek(0)
            with st.spinner("Loading files…"):
                loaded = ProjectLoader().load_many(project_files)
            if loaded and loaded.get("files"):
                st.session_state.project_data = loaded; st.session_state.sanitized_data = None
                st.session_state.result = None; st.session_state.prev_names = names
                st.success(f"✅ {loaded.get('summary','')}")

    if resume_file:
        resume_file.seek(0); raw = resume_file.read(); resume_file.seek(0)
        txt = extract_resume_text(resume_file)
        if txt and txt != st.session_state.resume_text:
            ext = pathlib.Path(resume_file.name).suffix.lower()
            st.session_state.resume_text = txt; st.session_state.resume_bytes = raw
            st.session_state.resume_ext = ext; st.session_state.resume_filename = pathlib.Path(resume_file.name).stem
            st.session_state.resume_profile = {}; st.session_state.result = None
            st.success(f"✅ Resume loaded ({ext.upper()}).")

    if st.session_state.project_data:
        info = st.session_state.project_data
        _gap(8)
        m1,m2,m3 = st.columns(3)
        m1.metric("Uploaded",info.get("uploaded_count",0)); m2.metric("Extracted",info.get("file_count",0)); m3.metric("Type",info.get("type","?").upper())

    _gap(14)
    _label("Actions")
    a1,a2,a3,a4 = st.columns(4)
    with a1:
        if st.button("🧹 Sanitize", key="btn_sanitize", use_container_width=True, disabled=not st.session_state.project_data):
            san = DataSanitizer(); raw2 = st.session_state.project_data.get("files",{})
            sd = dict(st.session_state.project_data); sd["files"] = {k: san.sanitize_text(v) for k,v in raw2.items()}
            st.session_state.sanitized_data = sd; st.session_state.san_report = san.get_report()
            st.success(f"✅ {st.session_state.san_report['redacted_values']} value(s) redacted.")
    with a2:
        if st.button("👤 Parse Resume", key="btn_parse_resume", use_container_width=True, disabled=not (api_key and st.session_state.resume_text)):
            with st.spinner("Parsing…"):
                try:
                    st.session_state.resume_profile = ResumeParser(provider,api_key,model).parse(st.session_state.resume_text)
                    st.success("✅ Profile extracted.")
                except Exception as e: st.error(f"Parse failed: {e}")
    with a3:
        ok = bool(api_key) and (st.session_state.sanitized_data or st.session_state.project_data) and st.session_state.resume_text
        if st.button("⚡ Generate Score", key="btn_generate_score", use_container_width=True, disabled=not ok, type="primary"):
            data  = st.session_state.sanitized_data or st.session_state.project_data
            ptype = st.session_state.get("pu_type","Software Development")
            pdesc = st.session_state.get("pu_desc","")
            pname = (st.session_state.get("pu_name","") or "").strip() or data.get("source","My Project")
            if not st.session_state.resume_profile:
                with st.spinner("Auto-parsing resume…"):
                    try: st.session_state.resume_profile = ResumeParser(provider,api_key,model).parse(st.session_state.resume_text)
                    except Exception: st.session_state.resume_profile={}
            with st.spinner(f"Scoring as {ptype} with {provider}…"):
                try:
                    st.session_state.result = TrustScorer(provider,api_key,model).score(data, st.session_state.resume_profile, project_type=ptype, description=pdesc)
                    st.session_state.result["source"] = pname; st.session_state.result["project_type"] = ptype; st.session_state.result["description"] = pdesc
                    save_submission(project_name=pname, result=st.session_state.result, profile=st.session_state.resume_profile,
                                    resume_text=st.session_state.resume_text, resume_bytes=st.session_state.resume_bytes or b"",
                                    resume_ext=st.session_state.resume_ext, resume_filename=st.session_state.resume_filename,
                                    uploaded_by=st.session_state.get("username","user"), project_type=ptype, description=pdesc)
                    st.success(f"✅ Scored as **{ptype}** and published!")
                    st.session_state.nav = "Results"; st.rerun()
                except Exception as e: st.error(f"Scoring failed: {e}")
    with a4:
        if st.button("🗑️ Clear All", key="btn_clear_all", use_container_width=True):
            for k in ["project_data","sanitized_data","resume_text","resume_profile","result","prev_names","san_report","resume_bytes","resume_ext","resume_filename"]:
                st.session_state.pop(k,None)
            st.rerun()
    _Wend()


# ══════════════════════════════════════════════════════════
# USER — RESULTS
# ══════════════════════════════════════════════════════════
def render_results():
    _page_header("📈","Trust Report","Full AI assessment of your project.")
    result = st.session_state.get("result"); profile = st.session_state.get("resume_profile",{}); resume_text = st.session_state.get("resume_text","")
    _W()
    b1,b2,b3,_ = st.columns([1,1,1,4])
    with b1:
        if st.button("🏠 Dashboard", key="rs_dash", use_container_width=True): st.session_state.nav="Dashboard"; st.rerun()
    with b2:
        if st.button("← Upload", key="rs_upload", use_container_width=True): st.session_state.nav="Upload"; st.rerun()
    with b3:
        if st.button("🗂️ Files", key="rs_files", use_container_width=True): st.session_state.nav="Files"; st.rerun()
    _gap(14)
    if not result:
        _gap(36)
        _,mc,_ = st.columns([1,2,1])
        with mc:
            st.info("No results yet — score a project first.")
            if st.button("→ Upload & Score", key="rs_goto_upload", type="primary", use_container_width=True): st.session_state.nav="Upload"; st.rerun()
        _Wend(); return
    rg = ReportGenerator()
    rg.render_candidate_profile(profile, resume_text, role="User")
    st.divider(); rg.render_trust_metrics(result); st.divider(); rg.render_summary_and_rec(result)
    st.divider()
    b1,b2,_ = st.columns([1,1,4])
    with b1:
        if st.button("← Score Again", key="rs_score_again", use_container_width=True): st.session_state.nav="Upload"; st.rerun()
    with b2:
        if st.button("🏠 Dashboard", key="rs_dash2", use_container_width=True): st.session_state.nav="Dashboard"; st.rerun()
    _Wend()


# ══════════════════════════════════════════════════════════
# USER — FILES
# ══════════════════════════════════════════════════════════
def render_files():
    _page_header("🗂️","File Manager","Browse, preview, and delete uploaded files.")
    _W()
    b1,b2,b3,_ = st.columns([1,1,1,4])
    with b1:
        if st.button("🏠 Dashboard", key="fm_dash", use_container_width=True): st.session_state.nav="Dashboard"; st.rerun()
    with b2:
        if st.button("← Upload", key="fm_upload", use_container_width=True): st.session_state.nav="Upload"; st.rerun()
    with b3:
        if st.session_state.get("result"):
            if st.button("📈 Results", key="fm_results", use_container_width=True): st.session_state.nav="Results"; st.rerun()
    _gap(14)
    display = st.session_state.get("sanitized_data") or st.session_state.get("project_data")
    if not display or not display.get("files"):
        st.info("No files uploaded yet."); _Wend(); return
    files = display.get("files",{}); fnames = sorted(files.keys())
    m1,m2,m3 = st.columns(3)
    m1.metric("Files",len(fnames)); m2.metric("Size",f"{sum(len(v) for v in files.values())//1024} KB"); m3.metric("Status","✅ Sanitized" if st.session_state.get("sanitized_data") else "⚠️ Raw")
    _gap(12)
    for k,dv in [("sel_file",fnames[0] if fnames else None),("del_file",None),("del_confirm",False)]:
        if k not in st.session_state: st.session_state[k]=dv
    cl,cr = st.columns([1,2.4],gap="large")
    with cl:
        _label("Files")
        for fname in fnames:
            ext=pathlib.Path(fname).suffix.lower(); short=pathlib.Path(fname).name[:28]
            ico=("🖼️" if ext in IMAGE_EXTENSIONS else "🎬" if ext in VIDEO_EXTENSIONS else "🐍" if ext==".py" else "📄")
            active=(st.session_state.sel_file==fname)
            bg=C["bluelt"] if active else C["s1"]; bdr=C["blue"] if active else C["bd"]; tc=C["blue"] if active else C["tx"]
            st.markdown(f"<div style='background:{bg};border:1px solid {bdr};border-radius:7px;padding:7px 9px;margin-bottom:3px;'><div style='font-size:12px;color:{tc};font-weight:{'600' if active else '400'};'>{ico} {short}</div></div>",unsafe_allow_html=True)
            rc1,rc2=st.columns([4,1])
            with rc1:
                if st.button("View",key=f"fm_v_{fname}",use_container_width=True):
                    st.session_state.sel_file=fname; st.session_state.del_file=None; st.session_state.del_confirm=False; st.rerun()
            with rc2:
                if st.button("🗑",key=f"fm_d_{fname}",use_container_width=True):
                    st.session_state.del_file=fname; st.session_state.del_confirm=True; st.rerun()
    with cr:
        if st.session_state.del_confirm and st.session_state.del_file:
            dname=st.session_state.del_file; st.warning(f"Delete `{pathlib.Path(dname).name}`?")
            y,n=st.columns(2)
            with y:
                if st.button("✅ Confirm Delete",key="fm_confirm_del",use_container_width=True):
                    for dk in ["project_data","sanitized_data"]:
                        d=st.session_state.get(dk)
                        if d: d.get("files",{}).pop(dname,None); d["file_count"]=len(d.get("files",{}))
                    remaining=[f for f in fnames if f!=dname]
                    st.session_state.sel_file=remaining[0] if remaining else None; st.session_state.del_file=None; st.session_state.del_confirm=False; st.rerun()
            with n:
                if st.button("❌ Cancel",key="fm_cancel_del",use_container_width=True):
                    st.session_state.del_file=None; st.session_state.del_confirm=False; st.rerun()
        sel=st.session_state.sel_file
        if sel and sel in files:
            _label(f"Preview — {pathlib.Path(sel).name}"); val=files[sel]
            if val.startswith("[IMAGE:"):
                try: st.image(base64.b64decode(val.split(":",2)[2].rstrip("]")),caption=sel,use_column_width=True)
                except: st.info("Image preview unavailable.")
            elif val.startswith("[VIDEO:"):
                rv=(st.session_state.get("project_data",{}).get("videos",{}).get(sel,{}).get("raw"))
                if rv: st.video(rv)
                else: st.info(f"📹 {val}")
            else:
                st.code(val[:5000],language=pathlib.Path(sel).suffix.lstrip("."or"text"))
    _Wend()


# ══════════════════════════════════════════════════════════
# ORG — PORTFOLIO
# ══════════════════════════════════════════════════════════
def render_portfolio():
    _page_header("🏆","Portfolio","All projects ranked by trust score.")
    _W()
    if st.button("🏠 Dashboard", key="pf_dash", use_container_width=False):
        st.session_state.nav="Dashboard"; st.rerun()
    _gap(12)
    real_subs=get_all_submissions(); all_projects=[]
    for sub in real_subs:
        r=sub.get("result",{})
        user_desc=sub.get("description","").strip() or r.get("executive_summary","")
        all_projects.append({"_type":"real","_key":sub.get("_key",""),"name":r.get("source",sub.get("_key","Project")),"desc":user_desc[:160],"project_type":sub.get("project_type",r.get("project_type","Software Development")),"score":r.get("trust_score",0),"grade":r.get("grade","?"),"working":r.get("working_percentage",0),"rec":r.get("recommendation","N/A"),"_sub":sub})
    for gh in GITHUB_PROJECTS:
        all_projects.append({"_type":"github","_key":gh["github"],"name":gh["name"],"desc":gh["description"][:160],"project_type":"Software Development","score":gh["score"],"grade":gh["grade"],"working":gh["working"],"rec":gh["recommendation"],"_gh":gh})
    all_projects.sort(key=lambda x:x["score"],reverse=True)
    if not all_projects: st.info("No projects yet."); _Wend(); return
    scores=[p["score"] for p in all_projects]
    s1,s2,s3,s4=st.columns(4)
    s1.metric("Total",len(all_projects)); s2.metric("Avg Score",f"{sum(scores)//len(scores)}/100"); s3.metric("Strong Hire",sum(1 for p in all_projects if p["rec"]=="Strong Hire")); s4.metric("Top Grade",all_projects[0]["grade"])
    _gap(14)
    hcols=st.columns([0.4,3.4,1.1,0.7,1.2,1.4,0.9])
    for col,txt in zip(hcols,["#","Project","Score","Grade","Working","Recommendation","Action"]):
        with col: st.markdown(f"<p style='font-size:10px;font-weight:700;color:{C['sub']};letter-spacing:.07em;text-transform:uppercase;padding-bottom:7px;border-bottom:2px solid {C['bd']};margin:0;'>{txt}</p>",unsafe_allow_html=True)
    _gap(4)
    for i,proj in enumerate(all_projects):
        gc=_grade_col(proj["grade"]); rc=_rec_col(proj["rec"])
        rank_s=("background:#f59e0b;color:#fff;" if i==0 else "background:#94a3b8;color:#fff;" if i==1 else "background:#b45309;color:#fff;" if i==2 else f"background:{C['s2']};color:{C['sub']};border:1px solid {C['bd']};")
        row=st.columns([0.4,3.4,1.1,0.7,1.2,1.4,0.9])
        with row[0]: st.markdown(f"<div style='width:24px;height:24px;border-radius:50%;{rank_s}display:flex;align-items:center;justify-content:center;font-size:10px;font-weight:700;margin-top:10px;'>{i+1}</div>",unsafe_allow_html=True)
        with row[1]:
            is_gh=proj["_type"]=="github"; ptype=proj.get("project_type",""); ptc=TYPE_COLOR.get(ptype,C["sub"])
            st.markdown(f"<div style='padding:3px 0;'><div style='display:flex;align-items:center;gap:5px;margin-bottom:2px;'><span style='font-size:12px;font-weight:600;color:{C['tx']};'>{'🐙' if is_gh else '📁'} {proj['name']}</span>"+(f"<span style='background:{ptc}15;color:{ptc};border:1px solid {ptc}30;border-radius:10px;padding:1px 7px;font-size:9px;font-weight:600;'>{ptype}</span>" if not is_gh else "")+"</div><div style='font-size:10px;color:{};line-height:1.4;'>{}</div></div>".format(C["sub"],proj["desc"][:80]+("…" if len(proj["desc"])>80 else "")),unsafe_allow_html=True)
        with row[2]: st.markdown(f"<div style='font-size:18px;font-weight:700;color:{C['blue']};font-family:monospace;margin-bottom:3px;'>{proj['score']}</div>",unsafe_allow_html=True); _progress(proj["score"],C["blue"],height=5)
        with row[3]: st.markdown(f"<div style='font-size:18px;font-weight:700;color:{gc};margin-top:4px;'>{proj['grade']}</div>",unsafe_allow_html=True)
        with row[4]: st.markdown(f"<div style='font-size:17px;font-weight:700;color:{C['tx']};margin-bottom:3px;'>{proj['working']}%</div>",unsafe_allow_html=True); _progress(proj["working"],C["green"],height=5)
        with row[5]: st.markdown(f"<div style='margin-top:5px;'>{_rec_badge(proj['rec'])}</div>",unsafe_allow_html=True)
        with row[6]:
            _gap(4)
            if st.button("View →",key=f"pf_view_{i}_{proj['_key']}",use_container_width=True):
                if proj["_type"]=="github": st.session_state.detail_gh=proj["_gh"]; st.session_state.detail_real=None
                else: st.session_state.detail_real=proj["_sub"]; st.session_state.detail_gh=None
                st.session_state.nav="Detail"; st.rerun()
        st.markdown(f"<div style='height:1px;background:{C['bd']};margin:4px 0;'></div>",unsafe_allow_html=True)
    _Wend()


# ══════════════════════════════════════════════════════════
# ORG — DETAIL
# ══════════════════════════════════════════════════════════
def render_detail():
    _page_header("📋","Project Detail","Full sanitized assessment.")
    _W()
    b1,b2,_=st.columns([1,1,5])
    with b1:
        if st.button("🏠 Dashboard",key="dt_dash",use_container_width=True): st.session_state.nav="Dashboard"; st.rerun()
    with b2:
        if st.button("← Portfolio",key="dt_port",use_container_width=True): st.session_state.nav="Portfolio"; st.rerun()
    _gap(10)
    gh=st.session_state.get("detail_gh"); real=st.session_state.get("detail_real")
    if not gh and not real: st.info("No project selected."); _Wend(); return
    if gh: _render_gh_detail(gh)
    else:  _render_real_detail(real)
    _Wend()


def _ip_notice(is_sample=False):
    label="IP-Protected Sample" if is_sample else "IP-Protected View"
    sub=("GitHub sample — details sanitized." if is_sample else "All metrics from sanitized AI analysis. No source code or credentials exposed.")
    st.markdown(f"<div style='background:{C['bluelt']};border:1px solid {C['bluebdr']};border-radius:8px;padding:10px 14px;margin-bottom:16px;display:flex;align-items:center;gap:9px;'><span style='font-size:16px;'>🔒</span><div><strong style='font-size:12px;color:{C['tx']};'>{label}</strong><p style='font-size:11px;color:{C['sub']};margin:1px 0 0;'>{sub}</p></div></div>",unsafe_allow_html=True)


def _render_gh_detail(gh):
    _ip_notice(is_sample=True); gc=_grade_col(gh["grade"]); rc=_rec_col(gh["recommendation"])
    t1,t2=st.columns([3,1])
    with t1: st.markdown(f"<h2 style='font-size:17px;font-weight:700;color:{C['tx']};margin-bottom:3px;'>🐙 {gh['name']}</h2><p style='font-size:12px;color:{C['sub']};margin-bottom:8px;'>github.com/{gh['github']}</p><p style='font-size:13px;color:{C['tx2']};line-height:1.7;'>{gh['description']}</p>",unsafe_allow_html=True)
    with t2: st.metric("Score",f"{gh['score']}/100"); st.metric("Grade",gh["grade"]); st.metric("Working",f"{gh['working']}%")
    st.divider(); _label("Tech Stack")
    st.markdown("<div style='display:flex;flex-wrap:wrap;gap:4px;margin-bottom:12px;'>"+"".join(_badge(t) for t in gh.get("tech_stack",[]))+"</div>",unsafe_allow_html=True)
    st.markdown(f"<div style='background:{C['s2']};border:1px solid {C['bd']};border-radius:7px;padding:12px;'><p style='font-size:13px;color:{C['tx2']};line-height:1.7;margin:0;'>📝 {gh.get('summary','')}</p></div>",unsafe_allow_html=True)
    st.divider(); st.button("✉️ Contact (Sample — N/A)",key="btn_contact_sample",disabled=True)


def _render_real_detail(sub):
    _ip_notice()
    result=sub.get("result",{}); profile=sub.get("profile",{}); resume_text=sub.get("resume_text","")
    ptype=sub.get("project_type",result.get("project_type","Software Development"))
    pdesc=sub.get("description",result.get("description","")).strip()
    ptc=TYPE_COLOR.get(ptype,C["sub"]); rc=_rec_col(result.get("recommendation","N/A"))
    name=profile.get("name","Unknown"); email=profile.get("email",""); source=result.get("source",sub.get("_key","Project"))
    t1,t2=st.columns([3,1])
    with t1:
        st.markdown(f"<h2 style='font-size:17px;font-weight:700;color:{C['tx']};margin-bottom:5px;'>📁 {source}</h2><div style='display:flex;align-items:center;gap:7px;margin-bottom:7px;'><span style='background:{ptc}15;color:{ptc};border:1px solid {ptc}30;border-radius:12px;padding:2px 10px;font-size:11px;font-weight:700;'>{ptype}</span><span style='font-size:11px;color:{C['sub']};'>by {sub.get('uploaded_by','')} · {sub.get('saved_at','')}</span></div>",unsafe_allow_html=True)
        if pdesc: st.markdown(f"<div style='background:{C['s2']};border-left:3px solid {ptc};border-radius:7px;padding:10px 13px;'><p style='font-size:13px;color:{C['tx2']};line-height:1.7;margin:0;'>{pdesc}</p></div>",unsafe_allow_html=True)
    with t2: st.metric("Score",f"{result.get('trust_score',0)}/100"); st.metric("Grade",result.get("grade","?")); st.metric("Working",f"{result.get('working_percentage',0)}%")
    st.divider(); _label("Candidate Profile")
    name_=profile.get("name","Unknown"); title_=profile.get("title",""); email_=profile.get("email","")
    phone_=profile.get("phone",""); loc_=profile.get("location",""); skills_=profile.get("top_skills",[])
    edu_=profile.get("education",""); bio_=profile.get("summary_bio",""); exp_=profile.get("years_experience","")
    linkedin_=profile.get("linkedin",""); github_=profile.get("github","")
    meta=" · ".join(filter(None,[title_,f"📍 {loc_}" if loc_ else "",f"⏱ {exp_} yrs" if exp_ else ""]))
    contacts="".join(filter(None,[f"<span style='margin-right:12px;'>📧 {email_}</span>" if email_ else "",f"<span style='margin-right:12px;'>📱 {phone_}</span>" if phone_ else "",f"<a href='{linkedin_}' target='_blank' style='color:{C['blue']};margin-right:12px;'>🔗 LinkedIn</a>" if linkedin_ else "",f"<a href='{github_}' target='_blank' style='color:{C['blue']};margin-right:12px;'>🐙 GitHub</a>" if github_ else ""]))
    chips="".join(f"<span style='background:{C['bluelt']};color:{C['blue']};border:1px solid {C['bluebdr']};padding:3px 10px;border-radius:5px;font-size:12px;font-weight:500;display:inline-block;margin:3px 3px 3px 0;'>{s}</span>" for s in skills_)
    st.markdown(f"<div style='background:{C['s1']};border:1px solid {C['bd']};border-radius:10px;padding:18px 20px;box-shadow:0 1px 3px rgba(0,0,0,.07);margin-bottom:10px;'><div style='font-size:17px;font-weight:700;color:{C['tx']};margin-bottom:2px;'>{name_}</div>"+(f"<div style='font-size:12px;color:{C['sub']};margin-bottom:8px;'>{meta}</div>" if meta else "")+(f"<div style='font-size:12px;color:{C['tx2']};margin-bottom:9px;padding-bottom:9px;border-bottom:1px solid {C['bd']};'>{contacts}</div>" if contacts else "")+(f"<div style='font-size:13px;color:{C['tx2']};background:{C['s2']};border-radius:6px;padding:9px 11px;margin-bottom:9px;line-height:1.7;'>{bio_}</div>" if bio_ else "")+(f"<div>{chips}</div>" if chips else "")+(f"<div style='font-size:12px;color:{C['sub']};margin-top:7px;'>🎓 {edu_}</div>" if edu_ else "")+"</div>",unsafe_allow_html=True)
    r_bytes=None; b64d=sub.get("resume_bytes_b64")
    if b64d:
        try: r_bytes=base64.b64decode(b64d)
        except: pass
    safe_resume_download(resume_text,profile,is_recruiter=True,resume_bytes=r_bytes,resume_ext=sub.get("resume_ext",".txt"),resume_filename=sub.get("resume_filename","resume"))
    st.divider(); rg=ReportGenerator(); rg.render_trust_metrics(result); st.divider(); rg.render_summary_and_rec(result)
    st.divider(); _label("Contact Candidate")
    if email_:
        rec_user=st.session_state.get("username","org"); org_email=_all_credentials().get(rec_user,{}).get("email",f"{rec_user}@company.com")
        sc_=result.get("trust_score",0); gr_=result.get("grade","?"); rec_=result.get("recommendation","")
        st.markdown(f"<div style='background:{C['s2']};border:1px solid {C['bd']};border-radius:8px;padding:12px 14px;margin-bottom:12px;'><div style='display:flex;justify-content:space-between;flex-wrap:wrap;gap:10px;'><div><div style='font-size:13px;font-weight:600;color:{C['tx']};'>👤 {name_}</div><div style='font-size:12px;color:{C['sub']};'>📧 {email_}</div></div><div style='text-align:right;'><div style='font-size:11px;color:{C['sub']};'>Your address</div><div style='font-size:12px;font-weight:600;color:{C['tx']};'>{org_email}</div></div></div></div>",unsafe_allow_html=True)
        subject=f"Re: Your Project Submission — DevTrustProtocol"
        body=f"Hi {name_},\n\nI reviewed your project '{source}' on DevTrustProtocol.\n\nTrust Score: {sc_}/100 | Grade: {gr_} | {rec_}\nProject Type: {ptype}\n\nI'd love to connect. Are you available for a call this week?\n\nBest regards,"
        btn_col,_=st.columns([2,5])
        with btn_col: _gmail_button(f"Open Gmail — Contact {name_.split()[0]}",email_,org_email,subject,body)
    else:
        _info_box("📧 No email found in resume.",C["amberlt"],C["amberbdr"],C["amber"])


# ══════════════════════════════════════════════════════════
# MAIN ROUTER
# ══════════════════════════════════════════════════════════
def main():
    st.session_state.dark_mode = False
    if "page" not in st.session_state:
        st.session_state.page = "home"

    _css()

    logged_in = st.session_state.get("logged_in", False)
    role      = st.session_state.get("role", "")
    page      = st.session_state.get("page", "home")

    # ── Sidebar ALWAYS visible on every page ─────────────────
    provider, api_key, model, nav = _sidebar()

    if not logged_in:
        if page == "login":
            render_login()
        elif page == "register":
            render_register()
        elif page == "privacy":
            render_privacy()
        else:
            render_home()
        return

    if role == "User":
        if   nav == "Dashboard": render_user_dashboard(st.session_state["username"])
        elif nav == "Upload":    render_upload(provider, api_key, model)
        elif nav == "Results":   render_results()
        elif nav == "Files":     render_files()
        else:                    render_user_dashboard(st.session_state["username"])
    elif role == "Recruiter":
        if   nav == "Dashboard": render_org_dashboard(st.session_state["username"])
        elif nav == "Portfolio": render_portfolio()
        elif nav == "Detail":    render_detail()
        else:                    render_org_dashboard(st.session_state["username"])
    else:
        st.error("Unknown role.")

if __name__ == "__main__" or True:
    main()
